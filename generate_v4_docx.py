from __future__ import annotations

from pathlib import Path

from docx import Document


SRC = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v3.8.docx")
DST = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")


CHANGE_TEXT = (
    "修正接口名称/字段/示例 JSON 格式、URL/Message 复制错误、参数编号与描述等问题；"
    "补充 v4.0 检查修订。"
)


SIMPLE_REPL = {
    "Version: 3.8": "Version: 4.0",
    "Version：3.8": "Version：4.0",
    "EQP_ InnerOuterBindingReport": "EQP_InnerOuterBindingReport",
    "EQP_ GetSetOrPcsInfo": "EQP_GetSetOrPcsInfo",
    "EQP_ IcCodeReport": "EQP_IcCodeReport",
    "EQP_ InnerGetOuterReport": "EQP_InnerGetOuterReport",
    "EAP_ DateTimeSyncCommand": "EAP_DateTimeSyncCommand",
    "Out_panel_ID": "OutPanelId",
    "In_panel_ID": "InPanelId",
    "OuPanelId": "OutPanelId",
    "EQPStatus": "EqpStatus",
    '"Run "': '"Run"',
    '" Result "': '"Result"',
    '" Y"': '"Y"',
    '" UserName "': '"UserName"',
    '" Password "': '"Password"',
    '" UserName"': '"UserName"',
    '" Password"': '"Password"',
    "Updata": "Update",
    "ture": "true",
    "fales": "false",
    "拆分拆分子任务接口": "拆分子任务接口",
    "REST:POST http://IP:Port/api/ EQP_DetectionInfoReport": "REST:POST http://IP:Port/api/EQP_DetectionInfoReport",
    "REST:POST http://IP:Port/api/ EQP_Reportcoding": "REST:POST http://IP:Port/api/EQP_Reportcoding",
    "REST:POST http://IP:Port/api/ EQP_AluminumBindingReport": "REST:POST http://IP:Port/api/EQP_AluminumBindingReport",
    "http://ip:port/api/ EAP_InitialDataRequest": "http://ip:port/api/EAP_InitialDataRequest",
    "http://ip:Port/api/ EAP_InitialDataRequest": "http://ip:Port/api/EAP_InitialDataRequest",
    "CompletionStatus\": \"Success\"": "CompletionStatus\": \"Finished\"",
    "3f1a9d4a-6d81-45b3-9ddd-a3b9b1f1c1d3": "20250107121135343",
}


def normalize_json_quotes(text: str) -> str:
    if "{" not in text and "}" not in text and "[" not in text and "]" not in text:
        return text
    return (
        text.replace("“", '"')
        .replace("”", '"')
        .replace("，", ",")
        .replace("：", ":")
    )


def apply_simple_replacements(text: str) -> str:
    new = normalize_json_quotes(text)
    for old, val in SIMPLE_REPL.items():
        new = new.replace(old, val)
    return new


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    paragraph.text = text


def update_all_paragraphs(doc: Document) -> None:
    for p in doc.paragraphs:
        set_paragraph_text(p, apply_simple_replacements(p.text))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_text(p, apply_simple_replacements(p.text))


def rewrite_cell_examples(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                new = txt

                if '"AccountList"' in txt and "Password" in txt:
                    new = """{
    "Code": "0000",
    "Success": true,
    "Msg": " ",
    "DateTime": "20241127150000",
    "Content": {
        "AccountList": [
            {
                "UserName": "user1",
                "Name": "某某",
                "Permission": 1,
                "Password": "psw1"
            }
        ]
    },
    "RequestId": "20250107121135343"
}"""

                if '"Message": "EQP_IcCodeReport"' in txt and '"Lot"' in txt and '"Codes"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_Reportcoding
{
    "From": "EQP",
    "Message": "EQP_Reportcoding",
    "DateTime": "20241127150000",
    "Content": {
        "Lot": "L0349224",
        "Project": "H02GM21AM1-00",
        "Panelcode": "509422402",
        "Time": "2023-01-01 12:43:44",
        "UserName": "2003323-张三",
        "EqpId": "DRL002",
        "Event": "正常",
        "PnlFace": "0",
        "Result": "Success",
        "Codes": [
            {
                "code": "000051",
                "positionx": "105.29449462890625",
                "positiony": "78.559661865234375",
                "type": "0",
                "codetype": "3"
            }
        ]
    },
    "RequestId": "20250107121135343"
}"""

                if '"Message": "EQP_DefectDataReport"' in txt and '"TaskId"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_TaskCompleteReport
{
    "From": "EQP",
    "Message": "EQP_TaskCompleteReport",
    "DateTime": "2025/11/20 15:00:00",
    "Content": {
        "EqpId": "EQ01",
        "TaskId": "T001",
        "OkCount": 50
    },
    "RequestId": "20250107121135343"
}"""

                if '"Message": "EQP_ReportedDrillProgramPath"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_ReportedDrillProgramPath
{
    "From": "EQP",
    "Message": "EQP_ReportedDrillProgramPath",
    "DateTime": "2026/04/16 12:30:00",
    "RequestId": "20250107121135343",
    "Content": {
        "EqpId": "DRI80001",
        "JobId": "Job1",
        "FileType": 1,
        "DrillProgramInfos": [
            {
                "PanelId": "PANEL00001",
                "Face": "CS",
                "DrillProgramPath": "\\\\EAP_SERVER\\\\DrillData\\\\PANEL00001\\\\CS\\\\DRILL_01.nc"
            }
        ]
    }
}"""

                if '"Message": "EQP_AluminumBindingReport"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_AluminumBindingReport
{
    "From": "EQP",
    "Message": "EQP_AluminumBindingReport",
    "DateTime": "20241127150000",
    "Content": {
        "EqpId": "Eqp1",
        "AluminiumId": "H02GM21AM1-00",
        "PanelIds": [
            "panel001",
            "panel002"
        ]
    },
    "RequestId": "20250107121135343"
}"""

                if '"Message": "EQP_BackDrillCompletionReport"' in txt and '"AxisBindings"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_BackDrillCompletionReport
{
    "From": "EQP",
    "Message": "EQP_BackDrillCompletionReport",
    "DateTime": "2026/04/16 13:00:00",
    "RequestId": "20250107121135343",
    "Content": {
        "EqpId": "DRI80001",
        "Face": "CS",
        "JobId": "Job1",
        "AxisBindings": [
            {
                "AxisCode": "1",
                "PanelId": "PNL000001"
            },
            {
                "AxisCode": "2",
                "PanelId": "PNL000002"
            }
        ],
        "CompletionStatus": "Finished"
    }
}"""

                if '"Message": "EQP_CTProcessDataReport"' in txt:
                    new = new.replace(
                        "REST:POST http://IP:Port/api/EQP_BackDrillCompletionReport",
                        "REST:POST http://IP:Port/api/EQP_CTProcessDataReport",
                    )

                if '"Message": "EAP_TaskDownload"' in txt:
                    new = new.replace(
                        "REST:POST http://IP:Port/api/EAP_JobDataDownload",
                        "REST:POST http://IP:Port/api/EAP_TaskDownload",
                    )
                    new = new.replace("20251229103045000", "20250107121135343")

                if '"Message": "EQP_InnerOuterBindingReport"' in txt and '"InPanelId"' in txt:
                    new = """REST:POST http://IP:Port/api/EQP_InnerOuterBindingReport
{
    "From": "EQP",
    "Message": "EQP_InnerOuterBindingReport",
    "DateTime": "2024/11/27 15:00:00",
    "Content": {
        "EqpId": "EQ01",
        "OutPanelId": "OutPanel001",
        "InPanelId": "InPanel001"
    },
    "RequestId": "20250107121135343"
}"""

                if '"OutPanelId": "OutPanel001"\n              "InPanelId"' in new:
                    new = new.replace(
                        '"OutPanelId": "OutPanel001"\n              "InPanelId"',
                        '"OutPanelId": "OutPanel001",\n              "InPanelId"',
                    )
                if '"OutPanelId": "OutPanel001",\n}' in new:
                    new = new.replace('"OutPanelId": "OutPanel001",\n}', '"OutPanelId": "OutPanel001"\n}')

                if '"Message": "EQP_PPValidityAsk "' in new:
                    new = new.replace('"Message": "EQP_PPValidityAsk "', '"Message": "EQP_PPValidityAsk"')
                    new = new.replace('"MaterialLot":"ma1",', '"MaterialLot": "ma1"')

                if '"LotId": "001"\n"}' in new:
                    new = new.replace('"LotId": "001"\n"}', '"LotId": "001"\n}')

                if '"BaseCopper"' in new:
                    new = new.replace('"DielectricThickness": "0.1",\n"BaseCopper"', '"DielectricThickness": "0.1",\n        "BaseCopper"')

                if new != txt:
                    cell.text = new


def final_normalize_cells(doc: Document) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                new = apply_simple_replacements(cell.text)
                if new != cell.text:
                    cell.text = new


def update_numbering_and_descriptions(doc: Document) -> None:
    replacements = {
        "返回值列": "返回值列表",
        "验证成功为true，失败则false": "验证成功为 true，失败则 false",
        "4.2.1 | PanelId | string | 产品序列码": "4.4.1 | PanelId | string | 产品序列码",
        "4,5 | JobId | string | 任务名称": "4.5 | JobId | string | 任务名称",
        " | PortId | string | 设备端口ID": "4.4 | PortId | string | 设备端口ID",
    }
    for idx, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text
                if "EAP_" in txt or "返回值列" in txt or "PortId" in txt:
                    for old, new in replacements.items():
                        txt = txt.replace(old, new)
                    if idx >= 75:
                        txt = txt.replace("调用接口来源（EQP）", "调用接口来源（EAP）")
                    if "EQP_CTProcessDataReport" in txt:
                        txt = txt.replace("背钻每趟做完后上报完工信息", "CT检查机制程数据报告")
                    if cell.text != txt:
                        cell.text = txt


def update_change_log(doc: Document) -> None:
    table = doc.tables[0]
    existing = ["|".join(cell.text for cell in row.cells) for row in table.rows]
    if not any("|4.0|" in row for row in existing):
        cells = table.add_row().cells
        cells[0].text = "2026-06-01"
        cells[1].text = "张涣化"
        cells[2].text = "4.0"
        cells[3].text = CHANGE_TEXT


def main() -> None:
    doc = Document(str(SRC))
    update_change_log(doc)
    update_all_paragraphs(doc)
    rewrite_cell_examples(doc)
    update_numbering_and_descriptions(doc)
    final_normalize_cells(doc)
    doc.save(str(DST))
    print(DST)


if __name__ == "__main__":
    main()
