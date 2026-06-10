from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.models import ApiInterface, ApiParameter, InterfaceDirection, ParameterKind
from app.services.word_export import export_word_document


def test_word_export_creates_docx_with_watermark(tmp_path: Path):
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-037",
        name="测试接口",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_Test",
        caller="EQP",
        provider="EAP",
        requirement="测试需求",
        scenario="测试场景",
        service_description="测试服务",
    )
    output = tmp_path / "spec.docx"

    export_word_document(
        output,
        [interface],
        {1: {"From": "EQP", "Message": "EQP_Test", "Content": {}}},
        {1: {"Code": "0000", "Success": True, "Content": {}}},
        watermark_text="厂商查看",
    )

    assert output.exists()
    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    header_text = "\n".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
    )
    assert "珠海超毅 EAP-EQP API 接口通讯规格书" in body_text
    assert "EQP-EAP-037 测试接口" in body_text
    assert "厂商查看" in header_text


def test_word_export_preserves_template_preface_and_replaces_old_interfaces(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_heading("原规格书标题", level=0)
    template.add_paragraph("这是原规格书已有内容。")
    old_table = template.add_table(rows=0, cols=4)
    _add_test_row(old_table, "需求说明", "旧接口", "旧接口", "旧接口")
    _add_test_row(old_table, "接口名称", "OLD_API", "OLD_API", "OLD_API")
    template.add_paragraph("旧接口后面的正文也需要替换")
    old_log = template.add_table(rows=0, cols=3)
    _add_test_row(old_log, "日志范例", "请求", "OLD_LOG")
    template.save(template_path)

    interface = ApiInterface(
        id=1,
        code="EAP-EQP-009",
        name="启动设备",
        direction=InterfaceDirection.EAP_TO_EQP,
        api_name="EAP_StartMachine",
        caller="EAP",
        provider="EQP",
        requirement="启动设备需求",
        scenario="EAP 下发启动指令",
        service_description="启动设备服务",
    )
    output = tmp_path / "spec_from_template.docx"

    export_word_document(
        output,
        [interface],
        {1: {"From": "EAP", "Message": "EAP_StartMachine", "Content": {}}},
        {1: {"Code": "0000", "Success": True, "Content": {}}},
        template_path=template_path,
    )

    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "原规格书标题" in body_text
    assert "这是原规格书已有内容。" in body_text
    assert "接口内容" in body_text
    assert "EAP-EQP-009 启动设备" in body_text
    assert "OLD_API" not in table_text
    assert "OLD_LOG" not in table_text
    assert "旧接口后面的正文也需要替换" not in body_text


def test_word_export_replaces_old_interface_section_headings(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("保留的前言")
    template.add_paragraph("三、 接口内容")
    template.add_paragraph("1. EQP -> EAP 接口")
    template.add_paragraph("EQP-EAP-001 旧接口")
    old_table = template.add_table(rows=0, cols=4)
    _add_test_row(old_table, "需求说明", "旧", "旧", "旧")
    _add_test_row(old_table, "接口名称", "OLD_API", "OLD_API", "OLD_API")
    old_log = template.add_table(rows=0, cols=3)
    _add_test_row(old_log, "日志范例", "请求", "OLD_LOG")
    template.save(template_path)
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-001",
        name="连线检查",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_AliveCheck",
        caller="EQP",
        provider="EAP",
    )
    output = tmp_path / "replaced_headings.docx"

    export_word_document(output, [interface], {1: {}}, {1: {}}, template_path=template_path)

    document = Document(output)
    non_empty = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    assert non_empty.count("三、 接口内容") == 1
    assert non_empty.count("1. EQP -> EAP 接口") == 1
    heading_index = non_empty.index("1. EQP -> EAP 接口")
    assert non_empty[heading_index + 1] == "EQP-EAP-001 连线检查"
    assert "保留的前言" in non_empty


def test_word_export_template_without_heading_styles_does_not_crash(tmp_path: Path):
    template_path = tmp_path / "template_without_heading_style.docx"
    template = Document()
    template.styles.element.remove(template.styles["Heading 1"].element)
    template.add_paragraph("原规格书内容")
    template.save(template_path)

    interface = ApiInterface(
        id=1,
        code="EAP-EQP-010",
        name="停止设备",
        direction=InterfaceDirection.EAP_TO_EQP,
        api_name="EAP_StopMachine",
        caller="EAP",
        provider="EQP",
    )
    output = tmp_path / "export_without_heading_style.docx"

    export_word_document(
        output,
        [interface],
        {1: {"Content": {}}},
        {1: {"Content": {}}},
        template_path=template_path,
    )

    assert output.exists()
    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "原规格书内容" in body_text
    assert "接口内容" in body_text


def test_word_export_includes_formal_parameter_tables_and_saved_log_examples(tmp_path: Path):
    interface = ApiInterface(
        id=1,
        code="EAP-EQP-001",
        name="初始化状态请求",
        direction=InterfaceDirection.EAP_TO_EQP,
        api_name="EAP_InitialDataRequest",
        caller="EAP",
        provider="EQP",
        request_log_example='REST:POST http://IP:Port/api/EAP_InitialDataRequest\n{"Content":{"IP":"127.0.0.1"}}',
        response_log_example='{"Code":"0000","Content":{"Result":true}}',
    )
    parameters = [
        ApiParameter(
            interface_id=1,
            kind=ParameterKind.REQUEST,
            sort_order=1,
            field_name="IP",
            data_type="string",
            required=True,
            is_array=False,
            example_value="127.0.0.1",
            description="IP地址",
        ),
        ApiParameter(
            interface_id=1,
            kind=ParameterKind.RESPONSE,
            sort_order=1,
            field_name="Result",
            data_type="bool",
            required=False,
            is_array=True,
            example_value="true",
            description="处理结果",
        ),
    ]
    output = tmp_path / "spec_with_tables.docx"

    export_word_document(
        output,
        [interface],
        {1: {"Content": {"IP": "127.0.0.1"}}},
        {1: {"Content": {"Result": True}}},
        parameters_by_interface={1: parameters},
    )

    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )

    assert "请求参数列表" in table_text
    assert "返回值列表" in table_text
    assert "日志范例" in table_text
    assert "REST:POST http://IP:Port/api/EAP_InitialDataRequest" in table_text
    assert '"IP":"127.0.0.1"' in table_text
    assert '"Result":true' in table_text
    assert "字段" in table_text
    assert "IP" in table_text
    assert "string" in table_text
    assert "必填" in table_text
    assert "非必填" in table_text
    assert "数组" in table_text
    assert "IP地址" in table_text
    assert "Result" in table_text
    assert "bool" in table_text
    assert "处理结果" in table_text
    assert any(len(table.columns) == 4 for table in document.tables)
    assert any(len(table.columns) == 3 for table in document.tables)


def test_word_export_uses_white_style_for_parameter_detail_rows(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("三、 接口内容")
    template.add_paragraph("EQP-EAP-001 连线检查")
    table = template.add_table(rows=15, cols=4)
    _add_shading(table.rows[5], "F2B000")
    _add_shading(table.rows[6], "F2B000")
    _set_test_row(table, 0, "需求说明", "旧", "旧", "旧")
    _set_test_row(table, 2, "接口名称", "OLD_API", "OLD_API", "OLD_API")
    _set_test_row(table, 5, "请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表")
    _set_test_row(table, 6, "序号", "字段", "类型", "描述")
    _set_test_row(table, 7, "1", "From", "string", "调用接口来源")
    log_table = template.add_table(rows=2, cols=3)
    _set_test_row(log_table, 0, "日志范例", "请求", "{}")
    template.save(template_path)
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-001",
        name="连线检查",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_AliveCheck",
        caller="EQP",
        provider="EAP",
    )
    output = tmp_path / "styled.docx"

    export_word_document(
        output,
        [interface],
        {1: {"Content": {}}},
        {1: {"Content": {}}},
        template_path=template_path,
    )

    document = Document(output)
    exported = next(table for table in document.tables if len(table.columns) == 4)
    assert _cell_fill(exported.rows[5].cells[0]) == "F2B000"
    assert _cell_fill(exported.rows[6].cells[0]) == "F2B000"
    assert _cell_fill(exported.rows[7].cells[0]) != "F2B000"


def test_word_export_uses_original_direction_heading_format_and_text(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("三、 接口内容")
    direction_heading = template.add_paragraph()
    direction_heading.paragraph_format.left_indent = Pt(18)
    run = direction_heading.add_run("1. EQP -> EAP 接口")
    run.bold = True
    run.font.size = Pt(14)
    template.add_paragraph("EQP-EAP-001 连线检查")
    main_table = template.add_table(rows=13, cols=4)
    _set_test_row(main_table, 0, "需求说明", "旧", "旧", "旧")
    _set_test_row(main_table, 2, "接口名称", "OLD_API", "OLD_API", "OLD_API")
    _set_test_row(main_table, 5, "请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表")
    _set_test_row(main_table, 6, "序号", "字段", "类型", "描述")
    _set_test_row(main_table, 7, "1", "From", "string", "调用接口来源")
    log_table = template.add_table(rows=2, cols=3)
    _set_test_row(log_table, 0, "日志范例", "请求", "{}")
    template.save(template_path)
    interface = ApiInterface(
        id=1,
        code="EAP-EQP-001",
        name="初始化状态请求",
        direction=InterfaceDirection.EAP_TO_EQP,
        api_name="EAP_InitialDataRequest",
        caller="EAP",
        provider="EQP",
    )
    output = tmp_path / "direction_heading.docx"

    export_word_document(
        output,
        [interface],
        {1: {"Content": {}}},
        {1: {"Content": {}}},
        template_path=template_path,
    )

    document = Document(output)
    headings = [paragraph for paragraph in document.paragraphs if "EAP -> EQP 接口" in paragraph.text]
    assert len(headings) == 1
    assert headings[0].text == "2. EAP -> EQP 接口"
    assert headings[0].paragraph_format.left_indent.twips == 360
    assert headings[0].runs[0].bold is True
    assert headings[0].runs[0].font.size.pt == 14


def test_word_export_places_first_interface_immediately_after_direction_heading(tmp_path: Path):
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-001",
        name="连线检查",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_AliveCheck",
        caller="EQP",
        provider="EAP",
    )
    output = tmp_path / "compact.docx"

    export_word_document(output, [interface], {1: {}}, {1: {}})

    document = Document(output)
    non_empty = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    heading_index = non_empty.index("1. EQP -> EAP 接口")
    assert non_empty[heading_index + 1] == "EQP-EAP-001 连线检查"


def test_word_export_preserves_nested_parameter_sequences(tmp_path: Path):
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-005",
        name="设备任务进展信息上报",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_EquipmentJobDataProcessReport",
        caller="EQP",
        provider="EAP",
    )
    parameters = [
        ApiParameter(
            id=1,
            interface_id=1,
            kind=ParameterKind.REQUEST,
            sort_order=1,
            field_name="Ext",
            data_type="ExtInfo",
            description="扩展信息",
            enum_options='{"sequence": "4.6"}',
        ),
        ApiParameter(
            id=2,
            interface_id=1,
            parent_id=1,
            kind=ParameterKind.REQUEST,
            sort_order=2,
            field_name="NgPanelList",
            data_type="List<NgPanel>",
            description="高压失败的Panel列表",
            enum_options='{"sequence": "4.6.1", "parent_sequence": "4.6"}',
        ),
        ApiParameter(
            id=3,
            interface_id=1,
            parent_id=2,
            kind=ParameterKind.REQUEST,
            sort_order=3,
            field_name="NgPanel",
            data_type="",
            description="NgPanel",
            enum_options='{"parent_sequence": "4.6.1", "is_group": true}',
        ),
        ApiParameter(
            id=4,
            interface_id=1,
            parent_id=2,
            kind=ParameterKind.REQUEST,
            sort_order=4,
            field_name="PanelId",
            data_type="string",
            description="产品序列码",
            enum_options='{"sequence": "4.6.1.1", "parent_sequence": "4.6.1"}',
        ),
    ]
    output = tmp_path / "nested.docx"

    export_word_document(
        output,
        [interface],
        {1: {}},
        {1: {}},
        parameters_by_interface={1: parameters},
    )

    document = Document(output)
    table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
    assert "4.6.1" in table_text
    assert "4.6.1.1" in table_text
    assert "NgPanelList" in table_text
    assert "PanelId" in table_text


def test_word_export_starts_toc_on_new_page_after_change_history(tmp_path: Path):
    template_path = tmp_path / "template.docx"
    template = Document()
    change_log = template.add_table(rows=1, cols=4)
    _set_test_row(change_log, 0, "2026-06-10", "张涣化", "4.5", "更新规格书内容")
    template.add_paragraph("三、 接口内容")
    template.add_paragraph("1. EQP -> EAP 接口")
    template.add_paragraph("EQP-EAP-001 旧接口")
    old_table = template.add_table(rows=0, cols=4)
    _add_test_row(old_table, "需求说明", "旧", "旧", "旧")
    _add_test_row(old_table, "接口名称", "OLD_API", "OLD_API", "OLD_API")
    old_log = template.add_table(rows=0, cols=3)
    _add_test_row(old_log, "日志范例", "请求", "OLD_LOG")
    _insert_toc_sdt_after_first_table(template, "目录")
    template.save(template_path)
    interface = ApiInterface(
        id=1,
        code="EQP-EAP-001",
        name="连线检查",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_AliveCheck",
        caller="EQP",
        provider="EAP",
    )
    output = tmp_path / "toc_page_break.docx"

    export_word_document(output, [interface], {1: {}}, {1: {}}, template_path=template_path)

    document = Document(output)
    children = list(document.element.body)
    toc_index = next(
        index
        for index, child in enumerate(children)
        if child.tag.endswith("sdt") and "目录" in "".join(node.text or "" for node in child.iter() if node.tag.endswith("t"))
    )
    previous = children[toc_index - 1]
    assert any(
        node.tag.endswith("br") and node.get(qn("w:type")) == "page"
        for node in previous.iter()
    )


def test_word_export_groups_by_code_prefix_and_appends_new_eap_to_eqp_interface(tmp_path: Path):
    existing = ApiInterface(
        id=1,
        code="EAP-EQP-001",
        name="初始化状态请求",
        direction=InterfaceDirection.EAP_TO_EQP,
        api_name="EAP_InitialDataRequest",
        caller="EAP",
        provider="EQP",
    )
    new_interface = ApiInterface(
        id=2,
        code="EAP-EQP-040",
        name="其他打码数据上报",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EAP_OtherCodeData",
        caller="EAP",
        provider="EQP",
    )
    eqp_interface = ApiInterface(
        id=3,
        code="EQP-EAP-001",
        name="连线检查",
        direction=InterfaceDirection.EQP_TO_EAP,
        api_name="EQP_AliveCheck",
        caller="EQP",
        provider="EAP",
    )
    output = tmp_path / "grouped.docx"

    export_word_document(
        output,
        [new_interface, eqp_interface, existing],
        {1: {}, 2: {}, 3: {}},
        {1: {}, 2: {}, 3: {}},
    )

    document = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    eqp_heading_index = body_text.index("1. EQP -> EAP 接口")
    eap_heading_index = body_text.index("2. EAP -> EQP 接口")
    eqp_item_index = body_text.index("EQP-EAP-001 连线检查")
    existing_eap_index = body_text.index("EAP-EQP-001 初始化状态请求")
    new_eap_index = body_text.index("EAP-EQP-040 其他打码数据上报")

    assert eqp_heading_index < eqp_item_index < eap_heading_index
    assert eap_heading_index < existing_eap_index < new_eap_index


def _add_test_row(table, *values: str) -> None:
    row = table.add_row()
    for index, value in enumerate(values):
        row.cells[index].text = value


def _set_test_row(table, row_index: int, *values: str) -> None:
    for index, value in enumerate(values):
        table.rows[row_index].cells[index].text = value


def _add_shading(row, fill: str) -> None:
    from docx.oxml import OxmlElement

    for cell in row.cells:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        cell._tc.get_or_add_tcPr().append(shading)


def _cell_fill(cell) -> str:
    tc_pr = cell._tc.tcPr
    shading = tc_pr.find(qn("w:shd")) if tc_pr is not None else None
    return shading.get(qn("w:fill")) if shading is not None else ""


def _insert_toc_sdt_after_first_table(document: Document, title: str) -> None:
    sdt = OxmlElement("w:sdt")
    sdt_content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = title
    run.append(text)
    paragraph.append(run)
    sdt_content.append(paragraph)
    sdt.append(sdt_content)
    body = document.element.body
    first_table_index = next(index for index, child in enumerate(body) if child.tag.endswith("tbl"))
    body.insert(first_table_index + 1, sdt)
