from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlmodel import Session, SQLModel, create_engine, select

from app.database import get_session
from app.main import app
from app.models import ApiInterface, ApiParameter, ParameterKind, SpecTemplate
from app.routers import imports


@pytest.fixture
def client_with_engine(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    monkeypatch.setattr(imports, "TEMPLATE_DIR", tmp_path / "templates")

    def override_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_session
    try:
        yield TestClient(app), engine
    finally:
        app.dependency_overrides.clear()


def test_upload_docx_spec_template_saves_file_and_metadata(client_with_engine):
    client, engine = client_with_engine

    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "珠海超毅 EAP-EQP API 接口通讯规格书4.0.docx",
                b"fake docx content",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert "导入成功" in response.text
    assert "珠海超毅 EAP-EQP API 接口通讯规格书4.0.docx" in response.text

    with Session(engine) as session:
        template = session.exec(select(SpecTemplate).order_by(SpecTemplate.created_at.desc())).first()

    assert template is not None
    assert template.original_filename == "珠海超毅 EAP-EQP API 接口通讯规格书4.0.docx"
    assert Path(template.stored_path).exists()


def test_upload_rejects_non_docx_file(client_with_engine):
    client, _ = client_with_engine

    response = client.post(
        "/imports/spec",
        files={"spec_file": ("readme.txt", b"text", "text/plain")},
    )

    assert response.status_code == 400
    assert "只支持上传 .docx Word 文件" in response.text


def test_home_page_shows_imported_template_status(client_with_engine):
    client, _ = client_with_engine
    client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                b"docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    response = client.get("/")

    assert response.status_code == 200
    assert "已导入" in response.text
    assert "原规格书.docx" in response.text


def test_upload_docx_parses_interface_basics_and_shows_result(client_with_engine, tmp_path):
    client, engine = client_with_engine
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("EQP-EAP-010 设备状态上报", level=2)
    document.add_paragraph("接口名称 EQP_StatusReport")
    document.add_heading("EAP-EQP-011 启动设备", level=2)
    document.add_paragraph("接口名称 EAP_StartMachine")
    document.save(docx_path)

    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert "本次解析结果" in response.text
    assert "解析接口总数" in response.text
    assert "EQP-EAP-010" in response.text
    assert "EAP-EQP-011" in response.text

    with Session(engine) as session:
        interfaces = session.exec(select(ApiInterface).order_by(ApiInterface.code)).all()

    assert len(interfaces) == 2
    assert interfaces[0].code == "EAP-EQP-011"
    assert interfaces[1].code == "EQP-EAP-010"


def test_upload_docx_updates_existing_interface_codes(client_with_engine, tmp_path):
    client, engine = client_with_engine
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("EQP-EAP-010 设备状态上报", level=2)
    document.add_paragraph("接口名称 EQP_StatusReport")
    document.save(docx_path)

    client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    updated_docx_path = tmp_path / "updated.docx"
    updated_document = Document()
    updated_document.add_heading("EQP-EAP-010 设备状态变更上报", level=2)
    updated_document.add_paragraph("接口名称 EQP_StatusChanged")
    updated_document.save(updated_docx_path)
    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                updated_docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    assert "已存在覆盖" in response.text

    with Session(engine) as session:
        interfaces = session.exec(select(ApiInterface)).all()

    assert len(interfaces) == 1
    assert interfaces[0].name == "设备状态变更上报"
    assert interfaces[0].api_name == "EQP_StatusChanged"


def test_upload_docx_saves_parsed_request_and_response_parameters(client_with_engine, tmp_path):
    client, engine = client_with_engine
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("EQP-EAP-010 设备状态上报", level=2)
    document.add_paragraph("接口名称 EQP_StatusReport")
    _add_parameter_table(document, "请求参数", ["LotId", "string", "是", "L001", "批次号"])
    _add_parameter_table(document, "响应参数", ["Result", "bool", "是", "true", "处理结果"])
    document.save(docx_path)

    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200

    with Session(engine) as session:
        interface = session.exec(select(ApiInterface)).one()
        parameters = session.exec(select(ApiParameter).order_by(ApiParameter.sort_order)).all()

    assert interface.code == "EQP-EAP-010"
    assert len(parameters) == 2
    assert parameters[0].kind == ParameterKind.REQUEST
    assert parameters[0].field_name == "LotId"
    assert parameters[1].kind == ParameterKind.RESPONSE
    assert parameters[1].field_name == "Result"


def test_upload_docx_saves_nested_parameter_parent_links(client_with_engine, tmp_path):
    client, engine = client_with_engine
    docx_path = tmp_path / "nested.docx"
    document = Document()
    document.add_heading("EQP-EAP-005 设备任务进展信息上报", level=2)
    table = document.add_table(rows=0, cols=4)
    for values in [
        ["接口名称", "EQP_EquipmentJobDataProcessReport", "", ""],
        ["请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表"],
        ["序号", "字段", "类型", "描述"],
        ["4", "Content", "object", "参数内容"],
        ["Content", "Content", "Content", "Content"],
        ["4.6", "Ext", "ExtInfo", "扩展信息"],
        ["ExtInfo（高压测试机生产结束时上报）", "", "", ""],
        ["4.6.1", "NgPanelList", "List<NgPanel>", "高压失败的Panel列表"],
        ["NgPanel", "", "", ""],
        ["4.6.1.1", "PanelId", "string", "产品序列码"],
    ]:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    with Session(engine) as session:
        parameters = {
            item.field_name: item
            for item in session.exec(select(ApiParameter).order_by(ApiParameter.sort_order)).all()
        }

    assert parameters["Ext"].parent_id is None
    assert parameters["NgPanelList"].parent_id == parameters["Ext"].id
    assert parameters["PanelId"].parent_id == parameters["NgPanelList"].id
    assert '"sequence": "4.6.1.1"' in parameters["PanelId"].enum_options


def test_upload_docx_replaces_existing_parameters_for_same_interface(client_with_engine, tmp_path):
    client, engine = client_with_engine
    first_docx = tmp_path / "first.docx"
    first_document = Document()
    first_document.add_heading("EQP-EAP-010 设备状态上报", level=2)
    first_document.add_paragraph("接口名称 EQP_StatusReport")
    _add_parameter_table(first_document, "请求参数", ["OldField", "string", "是", "old", "旧字段"])
    first_document.save(first_docx)
    second_docx = tmp_path / "second.docx"
    second_document = Document()
    second_document.add_heading("EQP-EAP-010 设备状态上报", level=2)
    second_document.add_paragraph("接口名称 EQP_StatusReport")
    _add_parameter_table(second_document, "请求参数", ["NewField", "string", "是", "new", "新字段"])
    second_document.save(second_docx)

    for path in [first_docx, second_docx]:
        client.post(
            "/imports/spec",
            files={
                "spec_file": (
                    "原规格书.docx",
                    path.read_bytes(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    with Session(engine) as session:
        parameters = session.exec(select(ApiParameter)).all()

    assert len(parameters) == 1
    assert parameters[0].field_name == "NewField"


def test_upload_docx_saves_log_examples(client_with_engine, tmp_path):
    client, engine = client_with_engine
    docx_path = tmp_path / "source.docx"
    document = Document()
    document.add_heading("EAP-EQP-010 初始化状态请求", level=2)
    info_table = document.add_table(rows=1, cols=4)
    for index, value in enumerate(["接口名称", "EAP_InitialDataRequest", "EAP_InitialDataRequest", "EAP_InitialDataRequest"]):
        info_table.rows[0].cells[index].text = value
    log_table = document.add_table(rows=0, cols=3)
    for values in [
        ["日志范例", "请求", 'REST:POST http://IP:Port/api/EAP_InitialDataRequest\n{"From":"EAP"}'],
        ["日志范例", "应答", '{"Code":"0000","Success":true}'],
    ]:
        row = log_table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    response = client.post(
        "/imports/spec",
        files={
            "spec_file": (
                "原规格书.docx",
                docx_path.read_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    with Session(engine) as session:
        interface = session.exec(select(ApiInterface)).one()

    assert "REST:POST http://IP:Port/api/EAP_InitialDataRequest" in interface.request_log_example
    assert '"Code":"0000"' in interface.response_log_example


def _add_parameter_table(document: Document, title: str, values: list[str]) -> None:
    document.add_paragraph(title)
    table = document.add_table(rows=1, cols=5)
    headers = ["字段名", "类型", "必填", "示例值", "说明"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    row = table.add_row()
    for index, value in enumerate(values):
        row.cells[index].text = value
