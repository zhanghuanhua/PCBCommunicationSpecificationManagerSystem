from pathlib import Path

from fastapi.testclient import TestClient
from docx import Document
from sqlmodel import Session, SQLModel, create_engine, select

from app.database import get_session
from app.main import app
from app.models import ApiInterface, ApiParameter, InterfaceDirection, ParameterKind, SpecTemplate, SpecVersion
from app.routers import exports as exports_router
from app.services.pdf_export import PdfConversionError


def _latest_export(pattern: str) -> Path:
    return max(Path("exports").glob(pattern), key=lambda path: path.stat().st_mtime)


def test_export_center_page_shows_format_and_watermark_options():
    client = TestClient(app)

    response = client.get("/exports")

    assert response.status_code == 200
    assert "导出中心" in response.text
    assert "Word + PDF" in response.text
    assert "添加水印" in response.text
    assert "修改人姓名" in response.text
    assert "修改内容" in response.text
    assert "window.setTimeout" in response.text
    assert "选择文件夹" in response.text
    assert "系统会记住本次选择的位置" in response.text


def test_export_center_prefills_last_output_dir(tmp_path, monkeypatch):
    saved_dir = tmp_path / "saved"
    saved_dir.mkdir()
    monkeypatch.setattr(exports_router, "_last_output_dir", lambda: str(saved_dir))
    client = TestClient(app)

    response = client.get("/exports")

    assert response.status_code == 200
    assert str(saved_dir) in response.text


def test_select_output_dir_saves_last_choice(tmp_path, monkeypatch):
    selected_dir = tmp_path / "selected"
    selected_dir.mkdir()
    captured = {}
    monkeypatch.setattr(exports_router, "_last_output_dir", lambda: "")
    monkeypatch.setattr(exports_router, "_choose_export_dir", lambda initial_dir="": selected_dir)
    monkeypatch.setattr(exports_router, "_save_last_output_dir", lambda path: captured.setdefault("path", path))
    client = TestClient(app)

    response = client.post("/exports/select-output-dir")

    assert response.status_code == 200
    assert response.json() == {"selected": True, "path": str(selected_dir)}
    assert captured["path"] == selected_dir


def test_export_center_creates_markdown_file():
    client = TestClient(app)
    before = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.md"))

    response = client.post(
        "/exports",
        data={
            "export_format": "markdown",
            "watermark_text": "厂商查看",
            "target_version": "4.0",
        },
    )

    assert response.status_code == 200
    assert "导出结果" in response.text
    created = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.md")) - before
    assert created
    output = created.pop()
    assert output.exists()
    assert "珠海超毅 EAP-EQP API 接口通讯规格书" in output.read_text(encoding="utf-8")
    assert str(output.resolve()) in response.text


def test_export_center_uses_imported_template_for_word(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_heading("原规格书标题", level=0)
    template.add_paragraph("原规格书已有章节")
    template.save(template_path)

    with Session(engine) as session:
        spec_version = SpecVersion(
            version="4.0",
            original_filename="原规格书.docx",
            template_path=str(template_path),
        )
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        spec_version_id = spec_version.id
        session.add(
            SpecTemplate(
                original_filename="原规格书.docx",
                stored_path=str(template_path),
                file_size=template_path.stat().st_size,
            )
        )
        session.commit()

    def override_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_session
    before = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.docx"))

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "word",
                "watermark_text": "厂商查看",
                    "spec_version_id": str(spec_version_id),
                "target_version": "4.0",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    created = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.docx")) - before
    assert created
    output = created.pop()
    assert output.exists()
    exported = Document(output)
    body_text = "\n".join(paragraph.text for paragraph in exported.paragraphs)
    assert "原规格书标题" in body_text
    assert "原规格书已有章节" in body_text
    assert "接口内容" in body_text


def test_export_result_page_provides_download_links():
    client = TestClient(app)
    before = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.md"))

    response = client.post(
        "/exports",
        data={"export_format": "markdown", "target_version": "4.0"},
    )

    assert response.status_code == 200
    created = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.md")) - before
    assert created
    output = created.pop()
    assert f"/exports/download/{output.name}" not in response.text
    assert f"/exports/open-folder/{output.name}" in response.text
    assert "导出成功" in response.text


def test_word_pdf_export_reports_pdf_failure(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        spec_version = SpecVersion(version="4.0")
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        spec_version_id = spec_version.id

    def override_session():
        with Session(engine) as session:
            yield session

    def fake_word_export(output_path, *args, **kwargs):
        Path(output_path).write_text("word ok", encoding="utf-8")

    def fake_pdf_export(*args, **kwargs):
        raise PdfConversionError("PDF conversion failed")

    app.dependency_overrides[get_session] = override_session
    monkeypatch.setattr(exports_router, "export_word_document", fake_word_export)
    monkeypatch.setattr(exports_router, "export_pdf_document", fake_pdf_export)

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "word_pdf",
                "spec_version_id": str(spec_version_id),
                "target_version": "4.0",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert "PDF conversion failed" in response.text
    assert "导出成功" not in response.text


def test_export_center_passes_change_history_fields(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        spec_version = SpecVersion(version="4.2")
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        spec_version_id = spec_version.id

    captured = {}

    def override_session():
        with Session(engine) as session:
            yield session

    def fake_word_export(output_path, *args, **kwargs):
        captured.update(kwargs)
        Path(output_path).write_text("word ok", encoding="utf-8")

    app.dependency_overrides[get_session] = override_session
    monkeypatch.setattr(exports_router, "export_word_document", fake_word_export)

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "word",
                "spec_version_id": str(spec_version_id),
                "target_version": "4.2",
                "change_author": "张涣化",
                "change_description": "修正接口参数层级并更新导出内容。",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert captured["change_author"] == "张涣化"
    assert captured["change_description"] == "修正接口参数层级并更新导出内容。"


def test_export_center_builds_change_history_from_added_interfaces(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        source = SpecVersion(version="4.1")
        target = SpecVersion(version="4.2", source_version_id=1)
        session.add(source)
        session.add(target)
        session.commit()
        session.refresh(source)
        session.refresh(target)
        target.source_version_id = source.id
        session.add(
            ApiInterface(
                spec_version_id=source.id,
                code="EQP-EAP-037",
                name="热熔机叠板结果上传",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_HotPressReport",
                caller="EQP",
                provider="EAP",
            )
        )
        session.add(
            ApiInterface(
                spec_version_id=target.id,
                code="EQP-EAP-037",
                name="热熔机叠板结果上传",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_HotPressReport",
                caller="EQP",
                provider="EAP",
            )
        )
        session.add(
            ApiInterface(
                spec_version_id=target.id,
                code="EQP-EAP-038",
                name="其他打码数据上报",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_OtherCodeData",
                caller="EQP",
                provider="EAP",
            )
        )
        session.commit()
        spec_version_id = target.id

    captured = {}

    def override_session():
        with Session(engine) as session:
            yield session

    def fake_word_export(output_path, *args, **kwargs):
        captured.update(kwargs)
        Path(output_path).write_text("word ok", encoding="utf-8")

    app.dependency_overrides[get_session] = override_session
    monkeypatch.setattr(exports_router, "export_word_document", fake_word_export)

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "word",
                "spec_version_id": str(spec_version_id),
                "target_version": "4.2",
                "change_author": "张涣化",
                "change_description": "用户填写内容",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert captured["change_description"] == "1. 新增EQP-EAP-038 其他打码数据上报"


def test_export_center_builds_change_history_from_template_missing_interface(tmp_path, monkeypatch):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    template_path = tmp_path / "template.docx"
    template = Document()
    template.add_paragraph("EQP-EAP-037 热熔机叠板结果上传")
    template.save(template_path)

    with Session(engine) as session:
        spec_version = SpecVersion(
            version="4.2",
            template_path=str(template_path),
        )
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        spec_version_id = spec_version.id
        session.add(
            SpecTemplate(
                original_filename="template.docx",
                stored_path=str(template_path),
                file_size=template_path.stat().st_size,
            )
        )
        session.add(
            ApiInterface(
                spec_version_id=spec_version.id,
                code="EQP-EAP-037",
                name="热熔机叠板结果上传",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_HotPressReport",
                caller="EQP",
                provider="EAP",
            )
        )
        session.add(
            ApiInterface(
                spec_version_id=spec_version.id,
                code="EQP-EAP-038",
                name="其他打码数据上报",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_OtherCodeData",
                caller="EQP",
                provider="EAP",
            )
        )
        session.commit()

    captured = {}

    def override_session():
        with Session(engine) as session:
            yield session

    def fake_word_export(output_path, *args, **kwargs):
        captured.update(kwargs)
        Path(output_path).write_text("word ok", encoding="utf-8")

    app.dependency_overrides[get_session] = override_session
    monkeypatch.setattr(exports_router, "export_word_document", fake_word_export)

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "word",
                "spec_version_id": str(spec_version_id),
                "target_version": "4.2",
                "change_author": "张涣化",
                "change_description": "用户填写内容",
            },
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    assert captured["change_description"] == "1. 新增EQP-EAP-038 其他打码数据上报"


def test_export_download_returns_generated_file():
    client = TestClient(app)
    export_dir = Path("exports")
    export_dir.mkdir(exist_ok=True)
    output = export_dir / "download-test.md"
    output.write_text("download ok", encoding="utf-8")

    response = client.get(f"/exports/download/{output.name}")

    assert response.status_code == 200
    assert response.text == "download ok"


def test_export_center_word_includes_saved_parameters(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        spec_version = SpecVersion(version="4.0")
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        spec_version_id = spec_version.id
        interface = ApiInterface(
            spec_version_id=spec_version.id,
            code="EQP-EAP-201",
            name="参数导出测试",
            direction=InterfaceDirection.EQP_TO_EAP,
            api_name="EQP_ParamExport",
            caller="EQP",
            provider="EAP",
        )
        session.add(interface)
        session.commit()
        session.refresh(interface)
        session.add(
            ApiParameter(
                interface_id=interface.id,
                kind=ParameterKind.REQUEST,
                sort_order=1,
                field_name="LotId",
                data_type="string",
                example_value="L001",
                description="批次号",
            )
        )
        session.commit()

    def override_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_session
    before = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.docx"))

    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={"export_format": "word", "spec_version_id": str(spec_version_id), "target_version": "4.0"},
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    created = set(Path("exports").glob("EAP-EQP接口通讯规格书_v*.docx")) - before
    assert created
    output = created.pop()
    exported = Document(output)
    table_text = "\n".join(
        cell.text
        for table in exported.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "LotId" in table_text


def test_export_as_new_version_saves_version_record(tmp_path):
    engine = create_engine(f"sqlite:///{tmp_path / 'test.db'}")
    SQLModel.metadata.create_all(engine)
    with Session(engine) as session:
        spec_version = SpecVersion(version="4.0")
        session.add(spec_version)
        session.commit()
        session.refresh(spec_version)
        session.add(
            ApiInterface(
                spec_version_id=spec_version.id,
                code="EQP-EAP-001",
                name="连线检查",
                direction=InterfaceDirection.EQP_TO_EAP,
                api_name="EQP_AliveCheck",
                caller="EQP",
                provider="EAP",
            )
        )
        session.commit()
        spec_version_id = spec_version.id

    def override_session():
        with Session(engine) as session:
            yield session

    app.dependency_overrides[get_session] = override_session
    try:
        client = TestClient(app)
        response = client.post(
            "/exports",
            data={
                "export_format": "markdown",
                "spec_version_id": str(spec_version_id),
                "target_version": "4.1",
            },
            follow_redirects=True,
        )
    finally:
        app.dependency_overrides.clear()

    assert response.status_code == 200
    with Session(engine) as session:
        versions = session.exec(select(SpecVersion).order_by(SpecVersion.version)).all()
        copied = session.exec(
            select(ApiInterface).where(ApiInterface.version == "4.1")
        ).one()

    assert [item.version for item in versions] == ["4.0", "4.1"]
    assert copied.code == "EQP-EAP-001"
