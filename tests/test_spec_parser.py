from pathlib import Path

from docx import Document

from app.models import ParameterKind
from app.services.spec_parser import parse_interface_basics_from_docx


def test_parse_interface_basics_from_docx_extracts_codes_names_and_api(tmp_path: Path):
    docx_path = tmp_path / "spec.docx"
    document = Document()
    document.add_heading("EQP-EAP-001 连线检查", level=2)
    document.add_paragraph("接口名称 EQP_AliveCheck")
    document.add_heading("EAP-EQP-002 启动设备", level=2)
    document.add_paragraph("API 名称 EAP_StartMachine")
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    assert len(result) == 2
    assert result[0].code == "EQP-EAP-001"
    assert result[0].name == "连线检查"
    assert result[0].api_name == "EQP_AliveCheck"
    assert result[0].caller == "EQP"
    assert result[0].provider == "EAP"
    assert result[1].code == "EAP-EQP-002"
    assert result[1].name == "启动设备"
    assert result[1].api_name == "EAP_StartMachine"
    assert result[1].caller == "EAP"
    assert result[1].provider == "EQP"


def test_parse_interface_basics_from_docx_extracts_request_and_response_parameters(tmp_path: Path):
    docx_path = tmp_path / "spec.docx"
    document = Document()
    document.add_heading("EQP-EAP-001 连线检查", level=2)
    document.add_paragraph("接口名称 EQP_AliveCheck")
    document.add_paragraph("请求参数")
    request_table = document.add_table(rows=1, cols=5)
    headers = ["字段名", "类型", "必填", "示例值", "说明"]
    for index, header in enumerate(headers):
        request_table.rows[0].cells[index].text = header
    request_row = request_table.add_row()
    request_values = ["LotId", "string", "是", "L001", "批次号"]
    for index, value in enumerate(request_values):
        request_row.cells[index].text = value
    document.add_paragraph("响应参数")
    response_table = document.add_table(rows=1, cols=5)
    for index, header in enumerate(headers):
        response_table.rows[0].cells[index].text = header
    response_row = response_table.add_row()
    response_values = ["Result", "bool", "是", "true", "处理结果"]
    for index, value in enumerate(response_values):
        response_row.cells[index].text = value
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    assert len(result) == 1
    assert len(result[0].parameters) == 2
    assert result[0].parameters[0].kind == ParameterKind.REQUEST
    assert result[0].parameters[0].field_name == "LotId"
    assert result[0].parameters[0].data_type == "string"
    assert result[0].parameters[0].required is True
    assert result[0].parameters[0].example_value == "L001"
    assert result[0].parameters[0].description == "批次号"
    assert result[0].parameters[1].kind == ParameterKind.RESPONSE
    assert result[0].parameters[1].field_name == "Result"


def test_parse_interface_table_sections_extracts_content_parameters(tmp_path: Path):
    docx_path = tmp_path / "spec_table_sections.docx"
    document = Document()
    document.add_heading("EAP-EQP-001 初始化状态请求", level=2)
    table = document.add_table(rows=0, cols=4)
    for values in [
        ["接口名称", "EAP_InitialDataRequest", "EAP_InitialDataRequest", "EAP_InitialDataRequest"],
        ["请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表"],
        ["序号", "字段", "类型", "描述"],
        ["1", "From", "string", "调用接口来源"],
        ["4", "Content", "object", "参数内容"],
        ["Content", "Content", "Content", "Content"],
        ["4.1", "EqpId", "string", "设备 ID"],
        ["返回值列表", "返回值列表", "返回值列表", "返回值列表"],
        ["序号", "字段", "类型", "描述"],
        ["1", "Code", "string", "结果代码"],
        ["5", "Content", "object", "参数内容"],
        ["Content", "Content", "Content", "Content"],
        ["5.1", "ControlMode", "string", "控制模式"],
    ]:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    assert len(result) == 1
    request_parameters = [item for item in result[0].parameters if item.kind == ParameterKind.REQUEST]
    response_parameters = [item for item in result[0].parameters if item.kind == ParameterKind.RESPONSE]
    assert [item.field_name for item in request_parameters] == ["EqpId"]
    assert [item.field_name for item in response_parameters] == ["ControlMode"]


def test_parse_interface_table_sections_preserves_nested_sequences_and_groups(tmp_path: Path):
    docx_path = tmp_path / "nested_sections.docx"
    document = Document()
    document.add_heading("EQP-EAP-005 设备任务进展信息上报", level=2)
    table = document.add_table(rows=0, cols=4)
    for values in [
        ["接口名称", "EQP_EquipmentJobDataProcessReport", "", ""],
        ["请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表"],
        ["序号", "字段", "类型", "描述"],
        ["4", "Content", "object", "参数内容"],
        ["Content", "Content", "Content", "Content"],
        ["4.6", "Ext", "ExtInfo", "扩展信息"],
        ["ExtInfo（高压测试机生产结束时上报）", "", "", ""],
        ["4.6.1", "NgPanelList", "List<NgPanel>", "高压失败的Panel列表"],
        ["NgPanel", "", "", ""],
        ["4.6.1.1", "PanelId", "string", "产品序列码"],
    ]:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    request_parameters = [item for item in result[0].parameters if item.kind == ParameterKind.REQUEST]
    assert [(item.field_name, item.sequence, item.parent_sequence, item.is_group) for item in request_parameters] == [
        ("Ext", "4.6", "4", False),
        ("ExtInfo（高压测试机生产结束时上报）", "", "4.6", True),
        ("NgPanelList", "4.6.1", "4.6", False),
        ("NgPanel", "", "4.6.1", True),
        ("PanelId", "4.6.1.1", "4.6.1", False),
    ]


def test_parse_interface_main_table_extracts_summary_fields_and_api_name(tmp_path: Path):
    docx_path = tmp_path / "spec_main_table.docx"
    document = Document()
    document.add_heading("EQP-EAP-002 设备状态上报", level=2)
    table = document.add_table(rows=0, cols=4)
    for values in [
        ["需求说明", "设备状态上报", "设备状态上报", "设备状态上报"],
        ["使用场景", "设备状态发生变化时，上报到 EAP 系统记录", "设备状态发生变化时，上报到 EAP 系统记录", "设备状态发生变化时，上报到 EAP 系统记录"],
        ["接口名称", "EQP_EquipmentCurrentStatus", "EQP_EquipmentCurrentStatus", "EQP_EquipmentCurrentStatus"],
        ["接口方式", "接口调用方", "接口提供方", "接口服务描述"],
        ["Web API", "EQP", "EAP", "设备状态上报"],
    ]:
        row = table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    assert len(result) == 1
    assert result[0].requirement == "设备状态上报"
    assert result[0].scenario == "设备状态发生变化时，上报到 EAP 系统记录"
    assert result[0].api_name == "EQP_EquipmentCurrentStatus"
    assert result[0].caller == "EQP"
    assert result[0].provider == "EAP"
    assert result[0].service_description == "设备状态上报"


def test_parse_interface_log_examples_from_log_table(tmp_path: Path):
    docx_path = tmp_path / "spec_log_examples.docx"
    document = Document()
    document.add_heading("EAP-EQP-001 初始化状态请求", level=2)
    info_table = document.add_table(rows=1, cols=4)
    values = ["接口名称", "EAP_InitialDataRequest", "EAP_InitialDataRequest", "EAP_InitialDataRequest"]
    for index, value in enumerate(values):
        info_table.rows[0].cells[index].text = value
    log_table = document.add_table(rows=0, cols=3)
    for values in [
        ["日志范例", "请求", 'REST:POST http://IP:Port/api/EAP_InitialDataRequest\n{"From":"EAP"}'],
        ["日志范例", "应答", '{"Code":"0000","Success":true}'],
    ]:
        row = log_table.add_row()
        for index, value in enumerate(values):
            row.cells[index].text = value
    document.save(docx_path)

    result = parse_interface_basics_from_docx(docx_path)

    assert len(result) == 1
    assert "REST:POST http://IP:Port/api/EAP_InitialDataRequest" in result[0].request_log_example
    assert '"From":"EAP"' in result[0].request_log_example
    assert '"Code":"0000"' in result[0].response_log_example
