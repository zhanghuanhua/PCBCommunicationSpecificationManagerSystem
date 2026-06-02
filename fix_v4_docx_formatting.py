from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")
OUT_DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0_更新版.docx")


def set_run_font(run, font_name: str, size_pt: int) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:cs"), font_name)


def apply_paragraph_font(paragraph, font_name: str, size_pt: int) -> None:
    if not paragraph.runs and paragraph.text:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt)


def set_cell_text_preserve_basic(cell, text: str) -> None:
    cell.text = text


def clean_repeated_labels(doc: Document) -> int:
    count = 0
    patterns = [
        (re.compile(r"返回值列表(?:表)+"), "返回值列表"),
        (re.compile(r"请求参数列表(?:表)+"), "请求参数列表"),
        (re.compile(r"参数列表(?:表)+"), "参数列表"),
        (re.compile(r"返回值列(?:表)+"), "返回值列表"),
    ]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new = text
                for pat, repl in patterns:
                    new = pat.sub(repl, new)
                if new != text:
                    set_cell_text_preserve_basic(cell, new)
                    count += 1
    return count


def is_example_table(table) -> bool:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return "日志范例" in text


def apply_table_fonts(doc: Document) -> tuple[int, int]:
    normal_tables = 0
    example_tables = 0
    for table in doc.tables:
        if is_example_table(table):
            font_name = "Cambria"
            size = 7
            example_tables += 1
        else:
            font_name = "Microsoft YaHei"
            size = 8
            normal_tables += 1
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_paragraph_font(paragraph, font_name, size)
    return normal_tables, example_tables


def main() -> None:
    doc = Document(str(DOCX))
    cleaned = clean_repeated_labels(doc)
    normal_tables, example_tables = apply_table_fonts(doc)
    try:
        doc.save(str(DOCX))
        saved = DOCX
    except PermissionError:
        doc.save(str(OUT_DOCX))
        saved = OUT_DOCX
    print(
        f"saved={saved}\n"
        f"cleaned_labels={cleaned}\n"
        f"normal_tables_microsoft_yahei_8={normal_tables}\n"
        f"example_tables_cambria_7={example_tables}"
    )


if __name__ == "__main__":
    main()
