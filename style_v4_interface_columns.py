from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")
OUT_DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0_接口列字体修正版.docx")


def set_run_font(run, font_name: str, size_pt: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_paragraph_font(paragraph, font_name: str, size_pt: float) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if not paragraph.runs and paragraph.text:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt)


def is_example_table(table) -> bool:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return "日志范例" in text


def is_interface_table(table) -> bool:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return ("接口方式" in text and "接口调用方" in text and "接口提供方" in text) or (
        "参数列表" in text and "序号" in text and "字段" in text and "类型" in text and "描述" in text
    )


def has_parameter_area_started(row_texts: list[str], in_area: bool) -> bool:
    joined = "|".join(row_texts)
    if "接口方式" in row_texts or "接口调用方" in row_texts or "接口提供方" in row_texts:
        return True
    if "参数列表" in joined or "请求参数列表" in joined or "返回值列表" in joined:
        return True
    if "序号" in row_texts and "字段" in row_texts and "类型" in row_texts:
        return True
    return in_area


def style_interface_table(table) -> None:
    in_area = False
    for row in table.rows:
        texts = [cell.text.strip() for cell in row.cells]
        in_area = has_parameter_area_started(texts, in_area)
        if not in_area:
            # Top requirement/use-case/name rows stay normal business text.
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_font(p, "Microsoft YaHei", 8)
            continue

        # Section divider rows like "参数列表", "返回值列表", "Content" are labels.
        non_empty = [t for t in texts if t]
        is_single_label_row = len(set(non_empty)) <= 1 and non_empty
        if is_single_label_row and non_empty[0] not in {"序号", "字段", "类型", "描述"}:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_paragraph_font(p, "Microsoft YaHei", 8)
            continue

        last_idx = len(row.cells) - 1
        for idx, cell in enumerate(row.cells):
            if idx < last_idx:
                font_name, size = "Cambria", 7
            else:
                font_name, size = "Microsoft YaHei", 8
            for p in cell.paragraphs:
                set_paragraph_font(p, font_name, size)


def main() -> None:
    doc = Document(str(DOCX))
    styled = 0
    skipped_examples = 0
    for table in doc.tables:
        if is_example_table(table):
            skipped_examples += 1
            continue
        if is_interface_table(table):
            style_interface_table(table)
            styled += 1
    try:
        doc.save(str(DOCX))
        saved = DOCX
    except PermissionError:
        doc.save(str(OUT_DOCX))
        saved = OUT_DOCX
    print(f"saved={saved}")
    print(f"styled_interface_tables={styled}")
    print(f"skipped_example_tables={skipped_examples}")


if __name__ == "__main__":
    main()
