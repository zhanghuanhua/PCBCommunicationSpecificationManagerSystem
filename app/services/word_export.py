import copy
import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.models import ApiInterface, ApiParameter, InterfaceDirection, ParameterKind
from app.services.watermark import add_text_watermark


TOC_TITLE = "\u76ee\u5f55"
INTERFACE_CONTENT_TITLE = "\u63a5\u53e3\u5185\u5bb9"


def export_word_document(
    output_path: Path,
    interfaces: list[ApiInterface],
    request_examples: dict[int, dict],
    response_examples: dict[int, dict],
    watermark_text: str = "",
    template_path: Path | None = None,
    parameters_by_interface: dict[int, list[ApiParameter]] | None = None,
    document_version: str | None = None,
    change_author: str = "",
    change_description: str = "",
) -> Path:
    document = Document(template_path) if template_path else Document()
    add_text_watermark(document, watermark_text)
    templates = _find_interface_templates(document)
    version = document_version or _version_from_interfaces(interfaces)

    if template_path:
        _replace_all_version_text(document, version)
        _append_change_history(document, version, change_author, change_description)
        _remove_existing_interface_content(document)
        _ensure_toc_starts_on_new_page(document)
    else:
        _add_heading(document, "珠海超毅 EAP-EQP API 接口通讯规格书", level=0)
        document.add_paragraph(f"Version: {version}")
        document.add_paragraph("本文档由接口管理系统自动生成。")
        document.add_page_break()

    _add_template_like_paragraph(document, templates.heading, "三、 接口内容")
    _append_direction(
        document,
        interfaces,
        InterfaceDirection.EQP_TO_EAP,
        "1. EQP -> EAP 接口",
        templates,
        request_examples,
        response_examples,
        parameters_by_interface or {},
        version,
    )
    _append_direction(
        document,
        interfaces,
        InterfaceDirection.EAP_TO_EQP,
        "2. EAP -> EQP 接口",
        templates,
        request_examples,
        response_examples,
        parameters_by_interface or {},
        version,
    )
    _replace_all_version_text(document, version)
    _compact_interface_start(document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


class _InterfaceTemplates:
    def __init__(
        self,
        heading: CT_P | None,
        direction_heading: CT_P | None,
        title: CT_P | None,
        main_table: CT_Tbl | None,
        log_table: CT_Tbl | None,
    ) -> None:
        self.heading = heading
        self.direction_heading = direction_heading
        self.title = title
        self.main_table = main_table
        self.log_table = log_table


def _append_direction(
    document: Document,
    interfaces: list[ApiInterface],
    direction: InterfaceDirection,
    heading: str,
    templates: _InterfaceTemplates,
    request_examples: dict[int, dict],
    response_examples: dict[int, dict],
    parameters_by_interface: dict[int, list[ApiParameter]],
    document_version: str,
) -> None:
    _add_template_like_paragraph(document, templates.direction_heading, heading)
    for item in _interfaces_for_direction(interfaces, direction):
        key = item.id or 0
        item.version = document_version
        _add_template_like_paragraph(document, templates.title, f"{item.code} {item.name}")
        main_table = _append_template_table(document, templates.main_table, 4)
        _fill_interface_table(main_table, item, parameters_by_interface.get(key, []))
        document.add_paragraph("")
        log_table = _append_template_table(document, templates.log_table, 3)
        _fill_log_table(log_table, item, request_examples.get(key, {}), response_examples.get(key, {}))


def _interfaces_for_direction(
    interfaces: list[ApiInterface], direction: InterfaceDirection
) -> list[ApiInterface]:
    return [
        item
        for item in sorted(interfaces, key=_interface_sort_key)
        if _effective_direction(item) == direction
    ]


def _effective_direction(interface: ApiInterface) -> InterfaceDirection:
    code = interface.code.upper()
    if code.startswith("EQP-EAP-"):
        return InterfaceDirection.EQP_TO_EAP
    if code.startswith("EAP-EQP-"):
        return InterfaceDirection.EAP_TO_EQP
    return interface.direction


def _version_from_interfaces(interfaces: list[ApiInterface]) -> str:
    for item in interfaces:
        if item.version:
            return item.version
    return "4.0"


def _interface_sort_key(interface: ApiInterface) -> tuple[int, str]:
    code = interface.code.upper()
    if code.startswith("EQP-EAP-") or code.startswith("EAP-EQP-"):
        return (0, code)
    return (1, f"{interface.created_at.isoformat()}-{interface.id or 0}")


def _find_interface_templates(document: Document) -> _InterfaceTemplates:
    blocks = list(_iter_document_blocks(document))
    heading = None
    direction_heading = None
    title = None
    main_table = None
    log_table = None

    for index, block in enumerate(blocks):
        if isinstance(block, Paragraph) and block.text.strip() == "三、 接口内容":
            heading = block._p
        if isinstance(block, Paragraph) and "EQP -> EAP" in block.text and "接口" in block.text:
            direction_heading = block._p
        if isinstance(block, Paragraph) and "EQP-EAP-001" in block.text:
            title = block._p
            for next_block in blocks[index + 1 :]:
                if isinstance(next_block, Table) and _is_interface_table(next_block):
                    main_table = next_block._tbl
                    break
            continue
        if main_table is not None and isinstance(block, Table) and _is_log_table(block):
            log_table = block._tbl
            break

    return _InterfaceTemplates(heading, direction_heading, title, main_table, log_table)


def _iter_document_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _append_change_history(
    document: Document,
    version: str,
    author: str,
    description: str,
) -> None:
    author = author.strip()
    description = description.strip()
    if not author or not description:
        return
    table = _find_change_history_table(document)
    if table is None:
        return
    today = datetime.now().strftime("%Y-%m-%d")
    existing = {
        (
            _cell_text(row.cells[0]) if len(row.cells) > 0 else "",
            _cell_text(row.cells[1]) if len(row.cells) > 1 else "",
            _cell_text(row.cells[2]) if len(row.cells) > 2 else "",
            _cell_text(row.cells[3]) if len(row.cells) > 3 else "",
        )
        for row in table.rows
    }
    row_values = (today, author, version, description)
    if row_values in existing:
        return
    row = table.add_row()
    for index, value in enumerate(row_values):
        if index < len(row.cells):
            _set_cell_text(row.cells[index], value)
    _format_change_history_row(row)


def _format_change_history_row(row) -> None:
    for index, cell in enumerate(row.cells):
        alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 3 else WD_ALIGN_PARAGRAPH.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = alignment
            for run in paragraph.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(8)


def _find_change_history_table(document: Document) -> Table | None:
    for table in document.tables:
        if len(table.columns) < 4 or not table.rows:
            continue
        first_row_text = [_cell_text(cell) for cell in table.rows[0].cells[:4]]
        compact = "".join(first_row_text).replace(" ", "")
        if "日期" in compact and "作者" in compact and "版本号" in compact and "变更内容" in compact:
            return table
    return None


def _cell_text(cell) -> str:
    return cell.text.strip()


def _remove_existing_interface_content(document: Document) -> None:
    body = document.element.body
    children = list(body)
    start_index = None
    for index, child in enumerate(children):
        if not child.tag.endswith("p"):
            continue
        text = Paragraph(child, document).text.strip()
        if INTERFACE_CONTENT_TITLE in text:
            start_index = index
            break
    for index, child in enumerate(children):
        if start_index is not None:
            break
        if child.tag.endswith("tbl") and _is_interface_table(Table(child, document)):
            start_index = _previous_interface_heading_index(children, index, document)
            break
    if start_index is None:
        return
    for child in children[start_index:]:
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def _ensure_toc_starts_on_new_page(document: Document) -> None:
    body = document.element.body
    children = list(body)
    for index, child in enumerate(children):
        if child.tag.endswith("sdt") and _is_toc_element(child):
            if index == 0:
                return
            previous = children[index - 1]
            if _has_page_break(previous) or _has_section_break(previous):
                return
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
            body.remove(paragraph._p)
            body.insert(index, paragraph._p)
            return


def _compact_interface_start(document: Document) -> None:
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() != "三、 接口内容":
            continue
        paragraph.paragraph_format.page_break_before = False
        paragraph.paragraph_format.space_before = None
        paragraph.paragraph_format.space_after = None
        _remove_empty_paragraphs_before(document, paragraph, limit=80)
        for next_paragraph in paragraphs[index + 1 : index + 4]:
            if next_paragraph.text.strip():
                next_paragraph.paragraph_format.page_break_before = False
                next_paragraph.paragraph_format.space_before = None
                break
        return


def _remove_empty_paragraphs_before(document: Document, paragraph: Paragraph, limit: int) -> None:
    body = document.element.body
    children = list(body)
    try:
        index = children.index(paragraph._p)
    except ValueError:
        return
    removed = 0
    for candidate in reversed(children[max(0, index - limit) : index]):
        if removed >= limit:
            break
        if not candidate.tag.endswith("p"):
            continue
        candidate_paragraph = Paragraph(candidate, document)
        if candidate_paragraph.text.strip():
            continue
        if _has_page_break(candidate) or _has_section_break(candidate):
            continue
        body.remove(candidate)
        removed += 1


def _has_page_break(element) -> bool:
    return any(
        child.tag.endswith("br") and child.get(qn("w:type")) == "page"
        for child in element.iter()
    )


def _has_section_break(element) -> bool:
    return any(child.tag.endswith("sectPr") for child in element.iter())


def _is_toc_element(element) -> bool:
    text = _element_text(element)
    instructions = "".join(
        getattr(child, "text", None) or ""
        for child in element.iter()
        if child.tag.endswith("instrText")
    )
    return TOC_TITLE in text or "TOC" in instructions.upper()


def _element_text(element) -> str:
    return "".join(getattr(child, "text", None) or "" for child in element.iter() if child.tag.endswith("t"))


def _replace_all_version_text(document: Document, version: str) -> None:
    for paragraph in _iter_all_paragraphs(document):
        _replace_version_in_paragraph(paragraph, version)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_version_in_paragraph(paragraph, version)


def _iter_all_paragraphs(document: Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            yield paragraph
        for paragraph in section.footer.paragraphs:
            yield paragraph


def _replace_version_in_paragraph(paragraph: Paragraph, version: str) -> None:
    text = paragraph.text
    if not text:
        return
    new_text = re.sub(r"(Version\s*:\s*)\d+(?:\.\d+)*", rf"\g<1>{version}", text, flags=re.IGNORECASE)
    new_text = re.sub(r"([vV])\d+(?:\.\d+)*", rf"\g<1>{version}", new_text)
    if new_text != text:
        _set_paragraph_text(paragraph, new_text)


def _previous_interface_heading_index(children, table_index: int, document: Document) -> int:
    for index in range(table_index - 1, -1, -1):
        child = children[index]
        if not child.tag.endswith("p"):
            continue
        text = Paragraph(child, document).text.strip()
        if "EQP-EAP-" in text or "EAP-EQP-" in text:
            return index
        if "接口内容" in text:
            return index
    return table_index


def _add_template_like_paragraph(document: Document, template: CT_P | None, text: str) -> Paragraph:
    if template is None:
        paragraph = document.add_paragraph(text)
        return paragraph
    element = copy.deepcopy(template)
    document.element.body.append(element)
    paragraph = Paragraph(element, document)
    _set_paragraph_text(paragraph, text)
    _clear_page_break_before(paragraph)
    return paragraph


def _clear_page_break_before(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.page_break_before = False
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = p_pr.find(qn("w:pageBreakBefore"))
    if page_break is not None:
        p_pr.remove(page_break)


def _append_template_table(document: Document, template: CT_Tbl | None, cols: int) -> Table:
    if template is None:
        return document.add_table(rows=0, cols=cols)
    element = copy.deepcopy(template)
    document.element.body.append(element)
    return Table(element, document)


def _fill_interface_table(table: Table, interface: ApiInterface, parameters: list[ApiParameter]) -> None:
    _ensure_table_rows(table, 13)
    _set_row(table, 0, ["需求说明", interface.requirement, interface.requirement, interface.requirement])
    _set_row(table, 1, ["使用场景", interface.scenario, interface.scenario, interface.scenario])
    _set_row(table, 2, ["接口名称", interface.api_name, interface.api_name, interface.api_name])
    _set_row(table, 3, ["接口方式", "接口调用方", "接口提供方", "接口服务描述"])
    _set_row(table, 4, ["Web API", interface.caller, interface.provider, interface.service_description])

    request_rows = _parameter_rows(parameters, ParameterKind.REQUEST)
    response_rows = _parameter_rows(parameters, ParameterKind.RESPONSE)
    rows = [
        ("section", ["请求参数列表", "请求参数列表", "请求参数列表", "请求参数列表"]),
        ("header", ["序号", "字段", "类型", "描述"]),
        ("detail", ["1", "From", "string", f"调用接口来源（{interface.caller}）"]),
        ("detail", ["2", "Message", "string", "消息（接口名）"]),
        ("detail", ["3", "DateTime", "DateTime", "系统时间（秒）（格式：yyyy/MM/dd HH:mm:ss）"]),
        ("detail", ["4", "Content", "object", "参数内容" if request_rows else "空"]),
        ("detail", ["5", "RequestId", "string", "请求 ID（17位唯一标识符：yyyyMMddHHmmssfff）（Format:毫秒时间格式）"]),
        ("merged", ["Content", "", "", ""]),
        *_typed_parameter_rows(request_rows),
        ("section", ["返回值列表", "返回值列表", "返回值列表", "返回值列表"]),
        ("header", ["序号", "字段", "类型", "描述"]),
        ("detail", ["1", "Code", "string", "结果代码(0000=成功， 其余为错误代号)"]),
        ("detail", ["2", "Success", "bool", "执行成功与否"]),
        ("detail", ["3", "Msg", "string", "提示讯息"]),
        ("detail", ["4", "DateTime", "DateTime", "系统时间（秒）（格式：yyyy/MM/dd HH:mm:ss）"]),
        ("detail", ["5", "Content", "object", "参数内容" if response_rows else "空"]),
        *_typed_parameter_rows(response_rows),
        ("detail", ["6", "RequestId", "string", "回复请求 ID（EAP返回的17位的唯一标识，与请求的RequestID一致）"]),
    ]
    _replace_rows_from(table, 5, rows)


def _parameter_rows(parameters: list[ApiParameter], kind: ParameterKind) -> list[list[str]]:
    rows = []
    items = [
        parameter
        for parameter in sorted(parameters, key=lambda item: (item.sort_order, item.id or 0))
        if parameter.kind == kind
    ]
    item_ids = {parameter.id for parameter in items if parameter.id is not None}
    top_level = [
        parameter
        for parameter in items
        if parameter.parent_id is None
        or (parameter.parent_id not in item_ids and not _parameter_parent_sequence(parameter))
    ]
    children_by_parent = _children_by_parent(items)
    for index, parameter in enumerate(top_level, start=1):
        rows.extend(_parameter_tree_rows(parameter, f"4.{index}", children_by_parent))
    return rows


def _typed_parameter_rows(parameter_rows: list[list[str]]) -> list[tuple[str, list[str]]]:
    typed_rows = []
    for row in parameter_rows:
        row_kind = "merged" if _is_merged_marker_row(row) else "detail"
        typed_rows.append((row_kind, row))
    return typed_rows


def _is_merged_marker_row(row: list[str]) -> bool:
    non_empty = [value for value in row if value]
    return len(non_empty) == 1 or len(set(non_empty)) == 1


def _children_by_parent(parameters: list[ApiParameter]) -> dict[int, list[ApiParameter]]:
    children: dict[int, list[ApiParameter]] = {}
    by_sequence = {
        sequence: parameter
        for parameter in parameters
        if (sequence := _parameter_sequence(parameter))
    }
    item_ids = {parameter.id for parameter in parameters if parameter.id is not None}
    for parameter in parameters:
        parent_id = parameter.parent_id if parameter.parent_id in item_ids else None
        parent_sequence = _parameter_parent_sequence(parameter)
        if parent_id is None and parent_sequence:
            parent = by_sequence.get(parent_sequence)
            parent_id = parent.id if parent and parent.id is not None else None
        if parent_id is None:
            continue
        children.setdefault(parent_id, []).append(parameter)
    for items in children.values():
        items.sort(key=lambda item: (item.sort_order, item.id or 0))
    return children


def _parameter_tree_rows(
    parameter: ApiParameter,
    fallback_sequence: str,
    children_by_parent: dict[int, list[ApiParameter]],
) -> list[list[str]]:
    if _is_group_parameter(parameter):
        rows = [[parameter.field_name, "", "", ""]]
        for index, child in enumerate(children_by_parent.get(parameter.id or 0, []), start=1):
            rows.extend(_parameter_tree_rows(child, f"{fallback_sequence}.{index}", children_by_parent))
        return rows
    sequence = _parameter_sequence(parameter) or fallback_sequence
    rows = [
        [
            sequence,
            parameter.field_name,
            parameter.data_type,
            _format_parameter_description(parameter),
        ]
    ]
    for index, child in enumerate(children_by_parent.get(parameter.id or 0, []), start=1):
        rows.extend(_parameter_tree_rows(child, f"{sequence}.{index}", children_by_parent))
    return rows


def _fill_log_table(table: Table, interface: ApiInterface, request_example: dict, response_example: dict) -> None:
    _ensure_table_rows(table, 2)
    request_text = interface.request_log_example or _format_request_log(interface, request_example)
    response_text = interface.response_log_example or json.dumps(response_example, ensure_ascii=False, indent=2)
    _set_row(table, 0, ["日志范例", "请求", request_text])
    _set_row(table, 1, ["日志范例", "应答", response_text])
    while len(table.rows) > 2:
        table._tbl.remove(table.rows[-1]._tr)


def _replace_rows_from(table: Table, start_index: int, rows: list[tuple[str, list[str]]]) -> None:
    section_template = copy.deepcopy(table.rows[start_index]._tr)
    header_template = copy.deepcopy(table.rows[start_index + 1]._tr if len(table.rows) > start_index + 1 else table.rows[start_index]._tr)
    detail_template = copy.deepcopy(table.rows[start_index + 2]._tr if len(table.rows) > start_index + 2 else table.rows[start_index]._tr)
    templates = {
        "section": section_template,
        "header": header_template,
        "detail": detail_template,
        "merged": detail_template,
    }
    while len(table.rows) > start_index:
        table._tbl.remove(table.rows[-1]._tr)
    for row_kind, row_values in rows:
        new_row = copy.deepcopy(templates[row_kind])
        table._tbl.append(new_row)
        current_row = table.rows[-1]
        if row_kind == "merged":
            _merge_row_cells(current_row)
            _set_merged_row_text(current_row, row_values[0] if row_values else "")
            _compact_row(current_row)
        else:
            _set_row(table, len(table.rows) - 1, row_values)


def _ensure_table_rows(table: Table, count: int) -> None:
    while len(table.rows) < count:
        table.add_row()


def _set_row(table: Table, row_index: int, values: list[str]) -> None:
    row = table.rows[row_index]
    for index, value in enumerate(values):
        if index < len(row.cells):
            _set_cell_text(row.cells[index], value)


def _merge_row_cells(row) -> None:
    if len(row.cells) <= 1:
        return
    first = row.cells[0]
    for cell in row.cells[1:]:
        first = first.merge(cell)


def _set_merged_row_text(row, text: str) -> None:
    if not row.cells:
        return
    cell = row.cells[0]
    _set_cell_text(cell, text)
    for paragraph in cell.paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def _compact_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for height in list(tr_pr.findall(qn("w:trHeight"))):
        tr_pr.remove(height)
    seen_cells = set()
    for cell in row.cells:
        if id(cell._tc) in seen_cells:
            continue
        seen_cells.add(id(cell._tc))
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1


def _set_cell_text(cell, text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        cell.text = text or ""
        return
    _set_paragraph_text(paragraphs[0], text or "")
    for paragraph in paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)


def _set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _format_request_log(interface: ApiInterface, request_example: dict) -> str:
    return f"REST:POST http://IP:Port/api/{interface.api_name}\n{json.dumps(request_example, ensure_ascii=False, indent=2)}"


def _format_parameter_description(parameter: ApiParameter) -> str:
    return parameter.description


def _parameter_sequence(parameter: ApiParameter) -> str:
    metadata = _parameter_metadata(parameter)
    sequence = metadata.get("sequence")
    return sequence if isinstance(sequence, str) else ""


def _parameter_metadata(parameter: ApiParameter) -> dict:
    if not parameter.enum_options:
        return {}
    try:
        data = json.loads(parameter.enum_options)
    except json.JSONDecodeError:
        return {}
    return data if isinstance(data, dict) else {}


def _parameter_parent_sequence(parameter: ApiParameter) -> str:
    metadata = _parameter_metadata(parameter)
    parent_sequence = metadata.get("parent_sequence")
    return parent_sequence if isinstance(parent_sequence, str) else ""


def _is_group_parameter(parameter: ApiParameter) -> bool:
    metadata = _parameter_metadata(parameter)
    return bool(metadata.get("is_group")) or not parameter.data_type.strip()


def _is_interface_table(table: Table) -> bool:
    if len(table.columns) != 4 or not table.rows:
        return False
    text = _table_text(table)
    return "需求说明" in text and "接口名称" in text


def _is_log_table(table: Table) -> bool:
    if len(table.columns) != 3:
        return False
    return "日志范例" in _table_text(table)


def _table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def _add_heading(document: Document, text: str, level: int):
    try:
        return document.add_heading(text, level=level)
    except KeyError:
        paragraph = document.add_paragraph(text)
        for run in paragraph.runs:
            run.bold = True
        return paragraph
