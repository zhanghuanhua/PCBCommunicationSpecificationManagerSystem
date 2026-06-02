from docx.document import Document as DocumentType


def add_text_watermark(document: DocumentType, watermark_text: str) -> None:
    if not watermark_text.strip():
        return

    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = watermark_text
