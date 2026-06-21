from datetime import datetime
import json
import os
from pathlib import Path
import re
import subprocess

from fastapi import APIRouter, Depends, Form, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from sqlmodel import Session, delete, select

from app.database import get_session
from app.models import ApiInterface, ApiParameter, ExportRecord, InterfaceStatus, SpecTemplate, SpecVersion
from app.services.examples import build_request_example, build_response_example
from app.services.markdown_export import render_markdown_document
from app.services.pdf_export import PdfConversionError, export_pdf_document
from app.services.word_export import export_word_document


router = APIRouter(prefix="/exports")
templates = Jinja2Templates(directory="app/templates")
EXPORT_DIR = Path("exports")
EXPORT_SETTINGS_PATH = Path("data/export_settings.json")


@router.get("")
def export_center(
    request: Request,
    spec_version_id: int | None = None,
    session: Session = Depends(get_session),
):
    spec_version = session.get(SpecVersion, spec_version_id) if spec_version_id else _latest_spec_version(session)
    return templates.TemplateResponse(
        request,
        "export_center.html",
        {
            "title": "导出中心",
            "spec_version": spec_version,
            "last_output_dir": _last_output_dir(),
        },
    )


@router.post("/select-output-dir")
def select_output_dir():
    selected = _choose_export_dir(_last_output_dir())
    if selected is None:
        return JSONResponse({"selected": False, "message": "未选择保存位置。"})
    _save_last_output_dir(selected)
    return {"selected": True, "path": str(selected)}


@router.post("")
def run_export(
    request: Request,
    export_format: str = Form(...),
    spec_version_id: int | None = Form(None),
    target_version: str = Form(""),
    change_author: str = Form(""),
    change_description: str = Form(""),
    watermark_enabled: bool = Form(False),
    watermark_text: str = Form(""),
    output_dir: str = Form(""),
    choose_output_dir: bool = Form(False),
    session: Session = Depends(get_session),
):
    spec_version = session.get(SpecVersion, spec_version_id) if spec_version_id else _latest_spec_version(session)
    if not spec_version:
        return templates.TemplateResponse(
            request,
            "export_center.html",
            {
                "title": "导出中心",
                "message": "暂无规格书版本，请先导入原规格书后再导出。",
                "spec_version": None,
            },
        )
    spec_version_id = spec_version.id or 0
    interfaces = session.exec(
        select(ApiInterface)
        .where(ApiInterface.spec_version_id == spec_version_id)
        .order_by(ApiInterface.code)
    ).all()
    parameters = session.exec(select(ApiParameter).order_by(ApiParameter.sort_order, ApiParameter.id)).all()
    parameters_by_interface = {
        item.id or 0: [parameter for parameter in parameters if parameter.interface_id == item.id]
        for item in interfaces
    }
    generated_at = datetime.now()
    request_examples = {
        item.id or 0: build_request_example(item, parameters_by_interface[item.id or 0], now=generated_at)
        for item in interfaces
    }
    response_examples = {
        item.id or 0: build_response_example(item, parameters_by_interface[item.id or 0], now=generated_at)
        for item in interfaces
    }
    output_files: list[str] = []
    watermark = watermark_text if watermark_enabled else ""
    export_dir = _selected_export_dir(output_dir, choose_output_dir)
    if export_dir is None:
        return templates.TemplateResponse(
            request,
            "export_center.html",
            {
                "title": "导出中心",
                "message": "未选择保存位置，已取消导出。",
                "spec_version": spec_version,
                "last_output_dir": _last_output_dir(),
            },
        )
    export_version = target_version.strip() or spec_version.version
    template = _template_for_spec(spec_version, session)
    template_path = Path(template.stored_path) if template else None
    change_author = change_author.strip()
    manual_change_description = change_description.strip()
    change_description = (
        _build_change_history_description(session, spec_version, template_path)
        or manual_change_description
        or _default_change_description(export_format)
    )
    export_name = f"EAP-EQP接口通讯规格书_v{export_version}_{datetime.now().strftime('%Y%m%d_%H%M%S_%f')}"

    if export_format in {"markdown", "all"}:
        markdown_path = export_dir / f"{export_name}.md"
        markdown_path.write_text(
            render_markdown_document(interfaces, request_examples, response_examples),
            encoding="utf-8",
        )
        output_files.append(str(markdown_path))

    word_path: Path | None = None
    if export_format in {"word", "word_pdf", "all", "pdf"}:
        word_path = export_dir / f"{export_name}.docx"
        export_word_document(
            word_path,
            interfaces,
            request_examples,
            response_examples,
            watermark,
            template_path=template_path,
            parameters_by_interface=parameters_by_interface,
            document_version=export_version,
            change_author=change_author,
            change_description=change_description,
        )
        if export_format in {"word", "word_pdf", "all"}:
            output_files.append(str(word_path))

    if export_format in {"pdf", "word_pdf", "all"}:
        pdf_path = export_dir / f"{export_name}.pdf"
        try:
            export_pdf_document(
                pdf_path,
                interfaces,
                request_examples,
                response_examples,
                parameters_by_interface,
                watermark,
                source_docx_path=word_path,
            )
            output_files.append(str(pdf_path))
        except PdfConversionError as exc:
            if export_format in {"pdf", "word_pdf", "all"}:
                return templates.TemplateResponse(
                    request,
                    "export_center.html",
                    {
                        "title": "导出中心",
                        "message": str(exc),
                        "spec_version": spec_version,
                    },
                )

    record = ExportRecord(
        version=export_version,
        scope="all",
        formats=export_format,
        watermark_enabled=watermark_enabled,
        watermark_text=watermark,
        output_path=";".join(str(Path(file).resolve()) for file in output_files),
        result="success",
    )
    session.add(record)
    _save_last_output_dir(export_dir)
    saved_spec_version = _save_exported_spec_version(
        session,
        source=spec_version,
        target_version=export_version,
        output_files=output_files,
        publish_interfaces=_exports_publish_interfaces(export_format),
    )
    if saved_spec_version.id == spec_version.id and _exports_publish_interfaces(export_format):
        _mark_version_interfaces_published(session, saved_spec_version.id or 0)
    session.commit()

    return templates.TemplateResponse(
        request,
        "export_result.html",
        {
            "title": "导出结果",
            "output_files": output_files,
            "spec_version": saved_spec_version,
            "exported_files": [
                {"name": Path(file).name, "absolute_path": str(Path(file).resolve())}
                for file in output_files
            ],
        },
    )


def _latest_spec_version(session: Session) -> SpecVersion | None:
    return session.exec(select(SpecVersion).order_by(SpecVersion.created_at.desc())).first()


def _template_for_spec(spec_version: SpecVersion, session: Session) -> SpecTemplate | None:
    if spec_version.template_path:
        template = session.exec(
            select(SpecTemplate).where(SpecTemplate.stored_path == spec_version.template_path)
        ).first()
        if template:
            return template
    return session.exec(select(SpecTemplate).order_by(SpecTemplate.created_at.desc())).first()


def _default_change_description(export_format: str) -> str:
    if export_format in {"word_pdf", "all"}:
        return "更新接口内容并导出 Word/PDF 文档。"
    if export_format == "word":
        return "更新接口内容并导出 Word 文档。"
    if export_format == "pdf":
        return "更新接口内容并导出 PDF 文档。"
    return "更新接口内容并导出审阅文档。"


def _build_change_history_description(
    session: Session,
    spec_version: SpecVersion,
    template_path: Path | None = None,
) -> str:
    current_interfaces = session.exec(
        select(ApiInterface).where(ApiInterface.spec_version_id == spec_version.id)
    ).all()
    current_by_code = {item.code: item for item in current_interfaces}
    changes: list[str] = []

    source = session.get(SpecVersion, spec_version.source_version_id) if spec_version.source_version_id else None
    if source:
        source_interfaces = session.exec(
            select(ApiInterface).where(ApiInterface.spec_version_id == source.id)
        ).all()
        source_by_code = {item.code: item for item in source_interfaces}
        for code in sorted(set(current_by_code) - set(source_by_code)):
            item = current_by_code[code]
            changes.append(f"新增{item.code} {item.name}")
        for code in sorted(set(source_by_code) - set(current_by_code)):
            item = source_by_code[code]
            changes.append(f"删除{item.code} {item.name}")
        for code in sorted(set(current_by_code) & set(source_by_code)):
            current = current_by_code[code]
            source_item = source_by_code[code]
            if _interface_changed(current, source_item):
                changes.append(f"更新{current.code} {current.name}")

    for item in _template_missing_new_interfaces(current_interfaces, template_path):
        change = f"新增{item.code} {item.name}"
        if change not in changes:
            changes.append(change)
    return "\n".join(f"{index}. {change}" for index, change in enumerate(changes, start=1))


def _template_missing_new_interfaces(
    current_interfaces: list[ApiInterface],
    template_path: Path | None,
) -> list[ApiInterface]:
    if not template_path or not template_path.exists():
        return []
    template_text = _template_document_text(template_path)
    if not template_text:
        return []
    max_numbers: dict[str, int] = {}
    for prefix, number in re.findall(r"\b((?:EQP-EAP|EAP-EQP)-)(\d{3})\b", template_text, flags=re.IGNORECASE):
        max_numbers[prefix.upper()] = max(max_numbers.get(prefix.upper(), 0), int(number))
    missing = []
    for item in current_interfaces:
        match = re.match(r"((?:EQP-EAP|EAP-EQP)-)(\d{3})$", item.code, flags=re.IGNORECASE)
        if not match:
            continue
        prefix = match.group(1).upper()
        number = int(match.group(2))
        if item.code not in template_text and number > max_numbers.get(prefix, 0):
            missing.append(item)
    return sorted(missing, key=lambda item: item.code)


def _template_document_text(template_path: Path) -> str:
    try:
        from docx import Document

        document = Document(template_path)
    except Exception:
        return ""
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return f"{paragraph_text}\n{table_text}"


def _interface_changed(current: ApiInterface, source: ApiInterface) -> bool:
    return any(
        getattr(current, field) != getattr(source, field)
        for field in [
            "name",
            "api_name",
            "caller",
            "provider",
            "requirement",
            "scenario",
            "service_description",
            "request_log_example",
            "response_log_example",
        ]
    )


def _save_exported_spec_version(
    session: Session,
    source: SpecVersion,
    target_version: str,
    output_files: list[str],
    publish_interfaces: bool = False,
) -> SpecVersion:
    if target_version == source.version:
        source.status = "EXPORTED"
        source.updated_at = datetime.now()
        session.add(source)
        return source

    existing = session.exec(select(SpecVersion).where(SpecVersion.version == target_version)).first()
    if existing:
        _replace_version_interfaces(session, source, existing, publish_interfaces)
        existing.status = "EXPORTED"
        existing.template_path = source.template_path
        existing.original_filename = Path(output_files[0]).name if output_files else source.original_filename
        existing.updated_at = datetime.now()
        session.add(existing)
        return existing

    target = SpecVersion(
        name=source.name,
        version=target_version,
        original_filename=Path(output_files[0]).name if output_files else source.original_filename,
        template_path=source.template_path,
        status="EXPORTED",
        source_version_id=source.id,
        updated_at=datetime.now(),
    )
    session.add(target)
    session.flush()
    _copy_version_interfaces(session, source, target, publish_interfaces)
    return target


def _replace_version_interfaces(
    session: Session,
    source: SpecVersion,
    target: SpecVersion,
    publish_interfaces: bool = False,
) -> None:
    existing_interfaces = session.exec(
        select(ApiInterface).where(ApiInterface.spec_version_id == target.id)
    ).all()
    for interface in existing_interfaces:
        session.exec(delete(ApiParameter).where(ApiParameter.interface_id == interface.id))
        session.delete(interface)
    session.flush()
    _copy_version_interfaces(session, source, target, publish_interfaces)


def _copy_version_interfaces(
    session: Session,
    source: SpecVersion,
    target: SpecVersion,
    publish_interfaces: bool = False,
) -> None:
    source_interfaces = session.exec(
        select(ApiInterface).where(ApiInterface.spec_version_id == source.id).order_by(ApiInterface.code)
    ).all()
    generated_at = datetime.now()
    for interface in source_interfaces:
        copied = ApiInterface(
            spec_version_id=target.id,
            code=interface.code,
            name=interface.name,
            direction=interface.direction,
            api_name=interface.api_name,
            method=interface.method,
            content_type=interface.content_type,
            caller=interface.caller,
            provider=interface.provider,
            requirement=interface.requirement,
            scenario=interface.scenario,
            service_description=interface.service_description,
            version=target.version,
            module=interface.module,
            status=InterfaceStatus.PUBLISHED if publish_interfaces else interface.status,
            remark=interface.remark,
            request_log_example=interface.request_log_example,
            response_log_example=interface.response_log_example,
            updated_at=datetime.now(),
        )
        session.add(copied)
        session.flush()
        parameters = session.exec(
            select(ApiParameter)
            .where(ApiParameter.interface_id == interface.id)
            .order_by(ApiParameter.sort_order, ApiParameter.id)
        ).all()
        old_to_new_parameter_id: dict[int, int] = {}
        copied_parameters: list[tuple[ApiParameter, ApiParameter]] = []
        for parameter in parameters:
            copied_parameter = ApiParameter(
                interface_id=copied.id or 0,
                kind=parameter.kind,
                parent_id=None,
                sort_order=parameter.sort_order,
                field_name=parameter.field_name,
                data_type=_normalize_data_type(parameter.data_type),
                required=parameter.required,
                is_array=parameter.is_array,
                example_value=parameter.example_value,
                description=parameter.description,
                enum_options=parameter.enum_options,
            )
            session.add(copied_parameter)
            session.flush()
            if parameter.id is not None and copied_parameter.id is not None:
                old_to_new_parameter_id[parameter.id] = copied_parameter.id
            copied_parameters.append((parameter, copied_parameter))
        for source_parameter, copied_parameter in copied_parameters:
            if source_parameter.parent_id is not None:
                copied_parameter.parent_id = old_to_new_parameter_id.get(source_parameter.parent_id)
                session.add(copied_parameter)
        session.flush()
        refreshed_parameters = [copied_parameter for _, copied_parameter in copied_parameters]
        copied.request_log_example = _format_request_log(
            copied,
            build_request_example(copied, refreshed_parameters, now=generated_at),
        )
        copied.response_log_example = json.dumps(
            build_response_example(copied, refreshed_parameters, now=generated_at),
            ensure_ascii=False,
            indent=4,
        )
        session.add(copied)


def _exports_publish_interfaces(export_format: str) -> bool:
    return export_format in {"word", "pdf", "word_pdf", "all"}


def _mark_version_interfaces_published(session: Session, spec_version_id: int) -> None:
    interfaces = session.exec(
        select(ApiInterface).where(ApiInterface.spec_version_id == spec_version_id)
    ).all()
    for interface in interfaces:
        interface.status = InterfaceStatus.PUBLISHED
        interface.updated_at = datetime.now()
        session.add(interface)


def _normalize_data_type(data_type: str) -> str:
    normalized = data_type.strip()
    if normalized.lower() == "int":
        return "int"
    return normalized


def _format_request_log(interface: ApiInterface, request_example: dict) -> str:
    return f"REST:POST http://IP:Port/api/{interface.api_name}\n{json.dumps(request_example, ensure_ascii=False, indent=4)}"


@router.get("/download/{filename}")
def download_export(filename: str):
    path = _resolve_export_file(filename)
    return FileResponse(path, filename=path.name)


@router.get("/open-folder/{filename}", response_class=HTMLResponse)
def open_export_folder(filename: str):
    path = _resolve_export_file(filename)
    folder = path.parent.resolve()
    opened = False
    try:
        os.startfile(folder)
        opened = True
    except OSError:
        opened = False
    return f"""
    <html>
      <body style="font-family: Microsoft YaHei, Arial, sans-serif; padding: 24px;">
        <h2>文件保存位置</h2>
        <p>{"已尝试打开文件夹。" if opened else "未能自动打开文件夹，请按以下路径手动打开。"}</p>
        <p>文件已保存在以下文件夹：</p>
        <p><code>{folder}</code></p>
        <p>文件名：<code>{path.name}</code></p>
        <p><a href="/exports/download/{path.name}">下载文件</a> | <a href="/exports">返回导出中心</a></p>
      </body>
    </html>
    """


def _resolve_export_file(filename: str) -> Path:
    export_dir = EXPORT_DIR.resolve()
    path = (export_dir / filename).resolve()
    if export_dir in path.parents and path.exists() and path.is_file():
        return path
    found = _find_recorded_export_file(filename)
    if found:
        return found
    raise HTTPException(status_code=404, detail="导出文件不存在")


def _find_recorded_export_file(filename: str) -> Path | None:
    from app.database import engine

    with Session(engine) as session:
        records = session.exec(select(ExportRecord).order_by(ExportRecord.created_at.desc()).limit(50)).all()
    for record in records:
        for value in record.output_path.split(";"):
            path = Path(value).resolve()
            if path.name == filename and path.exists() and path.is_file():
                return path
    return None


def _selected_export_dir(output_dir: str, choose_output_dir: bool = False) -> Path | None:
    selected = Path(output_dir).expanduser() if output_dir.strip() else None
    if selected is None and choose_output_dir:
        selected = _choose_export_dir(_last_output_dir())
        if selected is None:
            return None
    if not selected:
        selected = EXPORT_DIR
    selected.mkdir(parents=True, exist_ok=True)
    return selected


def _choose_export_dir(initial_dir: str = "") -> Path | None:
    selected = _choose_export_dir_with_tk(initial_dir)
    if selected is not None:
        return selected
    return _choose_export_dir_with_powershell(initial_dir)


def _choose_export_dir_with_tk(initial_dir: str = "") -> Path | None:
    try:
        import tkinter as tk
        from tkinter import filedialog

        root = tk.Tk()
        root.withdraw()
        root.attributes("-topmost", True)
        selected = filedialog.askdirectory(
            title="请选择规格书导出保存位置",
            initialdir=initial_dir if initial_dir and Path(initial_dir).exists() else None,
            mustexist=False,
        )
        root.destroy()
    except Exception:
        return None
    return Path(selected) if selected else None


def _choose_export_dir_with_powershell(initial_dir: str = "") -> Path | None:
    initial_dir = initial_dir.replace("'", "''") if initial_dir else ""
    script = (
        "Add-Type -AssemblyName System.Windows.Forms;"
        "$dialog = New-Object System.Windows.Forms.FolderBrowserDialog;"
        "$dialog.Description = '请选择规格书导出保存位置';"
        "$dialog.ShowNewFolderButton = $true;"
        f"if ('{initial_dir}' -and (Test-Path '{initial_dir}')) {{ $dialog.SelectedPath = '{initial_dir}'; }}"
        "if ($dialog.ShowDialog() -eq [System.Windows.Forms.DialogResult]::OK) {"
        "  [Console]::OutputEncoding = [System.Text.Encoding]::UTF8;"
        "  Write-Output $dialog.SelectedPath"
        "}"
    )
    try:
        result = subprocess.run(
            ["powershell", "-NoProfile", "-STA", "-Command", script],
            capture_output=True,
            text=True,
            timeout=120,
        )
    except (FileNotFoundError, subprocess.TimeoutExpired):
        return None
    if result.returncode != 0:
        return None
    selected = result.stdout.strip().splitlines()
    return Path(selected[-1]) if selected else None


def _last_output_dir() -> str:
    try:
        data = json.loads(EXPORT_SETTINGS_PATH.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return ""
    value = data.get("last_output_dir", "")
    return value if isinstance(value, str) else ""


def _save_last_output_dir(path: Path) -> None:
    EXPORT_SETTINGS_PATH.parent.mkdir(parents=True, exist_ok=True)
    EXPORT_SETTINGS_PATH.write_text(
        json.dumps({"last_output_dir": str(path.resolve())}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
