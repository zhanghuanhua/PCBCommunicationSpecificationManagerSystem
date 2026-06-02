from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")
FULL_WIDTH = 9360


def set_font(paragraph, name: str, size: float, bold: bool | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = name
        run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = r_pr._add_rFonts()
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            r_fonts.set(qn(f"w:{key}"), name)


def set_tbl_layout(table) -> None:
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")
    ind = tbl_pr.find(qn("w:tblInd"))
    if ind is not None:
        tbl_pr.remove(ind)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(FULL_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_grid(table, widths: list[int]) -> None:
    set_tbl_layout(table)
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def set_cell_margins(cell, top=28, start=55, bottom=28, end=55) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def is_example_table(table) -> bool:
    return "日志范例" in "\n".join(cell.text for row in table.rows for cell in row.cells)


def is_interface_table(table) -> bool:
    text = "\n".join(cell.text for row in table.rows for cell in row.cells)
    return "接口方式" in text and "接口调用方" in text and "接口提供方" in text


def apply_method_block_fonts(table) -> None:
    for row in table.rows:
        texts = [cell.text.strip() for cell in row.cells]
        if "接口调用方" in texts and "接口提供方" in texts:
            for idx, cell in enumerate(row.cells):
                if idx <= 2:
                    for p in cell.paragraphs:
                        set_font(p, "Cambria", 7)
                else:
                    for p in cell.paragraphs:
                        set_font(p, "Microsoft YaHei", 8)
        if len(texts) >= 4 and texts[0] == "Web API":
            for idx, cell in enumerate(row.cells):
                if idx <= 2:
                    for p in cell.paragraphs:
                        set_font(p, "Cambria", 7)
                else:
                    for p in cell.paragraphs:
                        set_font(p, "Microsoft YaHei", 8)


def keep_together_for_headings(doc: Document) -> None:
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith(("EQP-EAP-", "EAP-EQP-")):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)


def refine_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = False
        example = is_example_table(table)
        if example:
            widths = [1250, 1150, 6960]
            font, size = "Cambria", 7
            valign = WD_CELL_VERTICAL_ALIGNMENT.TOP
        elif len(table.columns) == 4:
            widths = [900, 1700, 1300, 5860]
            font, size = "Microsoft YaHei", 8
            valign = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif len(table.columns) == 5:
            widths = [650, 1250, 1250, 1250, 4660]
            font, size = "Microsoft YaHei", 8
            valign = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif len(table.columns) == 3:
            widths = [1050, 1700, 6610]
            font, size = "Microsoft YaHei", 8
            valign = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        elif len(table.columns) == 2:
            widths = [1400, 7960]
            font, size = "Microsoft YaHei", 8
            valign = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        else:
            widths = [FULL_WIDTH // len(table.columns)] * len(table.columns)
            font, size = "Microsoft YaHei", 8
            valign = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_grid(table, widths)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None and len(row.cells) <= 5:
                tr_pr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                set_cell_margins(cell)
                cell.vertical_alignment = valign
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_font(p, font, size)
        if example:
            # Keep the log label narrow but readable and request/response labels wide enough.
            for row in table.rows:
                for idx, cell in enumerate(row.cells[:2]):
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        set_font(p, "Cambria", 7, bold=(idx == 0))
        if is_interface_table(table):
            apply_method_block_fonts(table)


def remove_header_first_page_linking(doc: Document) -> None:
    for section in doc.sections:
        section.different_first_page_header_footer = False
        for header in (section.header, section.first_page_header, section.even_page_header):
            for p in header.paragraphs:
                if "Version:" in p.text:
                    p.text = p.text.replace("Version: 3.8", "Version: 4.0")
                    for run in p.runs:
                        run.font.size = Pt(9)


def clear_page_breaks_before_parameter_rows(docx_path: Path) -> None:
    # Remove empty manual page breaks that split a table-like interface block across pages.
    tmp = docx_path.with_suffix(".tmp.docx")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    from lxml import etree

    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                body = root.find("w:body", namespaces=ns)
                blocks = list(body)
                for i, block in enumerate(blocks):
                    has_break = bool(block.xpath(".//w:br[@w:type='page']", namespaces=ns))
                    text = "".join(block.xpath(".//w:t/text()", namespaces=ns)).strip()
                    if has_break and not text:
                        next_text = "".join(blocks[i + 1].xpath(".//w:t/text()", namespaces=ns)).strip() if i + 1 < len(blocks) else ""
                        if next_text.startswith(("1Fromstring", "序号字段类型描述", "Content")):
                            body.remove(block)
                data = etree.tostring(root, encoding="UTF-8", xml_declaration=True, standalone="yes")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main() -> None:
    doc = Document(str(DOCX))
    keep_together_for_headings(doc)
    remove_header_first_page_linking(doc)
    refine_tables(doc)
    doc.save(str(DOCX))
    clear_page_breaks_before_parameter_rows(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
