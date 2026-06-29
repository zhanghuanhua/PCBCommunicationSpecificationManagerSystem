from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
IMAGE_DIR = DOCS_DIR / "images" / "manual"
MD_PATH = DOCS_DIR / "EAP接口规格书管理系统操作手册.md"
DOCX_PATH = DOCS_DIR / "EAP接口规格书管理系统操作手册.docx"
TODAY = datetime.now().strftime("%Y-%m-%d")


SECTIONS = [
    {
        "level": 1,
        "title": "1. 系统概述",
        "paragraphs": [
            "EAP接口规格书管理系统用于集中维护 EAP 与 EQP 之间的 Web API 通讯规格书。系统支持导入原始 Word 规格书、维护接口与参数、生成日志范例，并导出 Word、PDF 或 Markdown 文档。",
            "系统部署到服务器后，开发人员通过浏览器访问同一个地址，接口数据、版本履历、导出记录和上传模板会统一保存在服务器上。",
        ],
        "bullets": [
            "适用人员：EAP 开发、EQP 对接开发、测试人员、规格书维护人员。",
            "推荐访问方式：http://服务器IP:8000。",
            "推荐维护原则：先确认当前规格书版本，再进入对应版本维护接口内容。",
        ],
    },
    {
        "level": 1,
        "title": "2. 规格书版本管理",
        "paragraphs": [
            "进入系统后默认显示规格书版本管理页面。页面按版本从新到旧排列，最新版本在最上方。每个版本卡片展示接口总数、接口方向统计和版本号。",
        ],
        "image": ("01-version-home.png", "图 1 规格书版本管理首页"),
        "steps": [
            "确认需要维护的规格书版本。",
            "点击“进入管理”进入该版本的接口管理页面。",
            "点击“导出”可直接进入该版本的导出中心。",
            "如需导入新的原始规格书，点击“导入原规格书”。",
        ],
        "notes": [
            "已经导出 Word 或 PDF 的版本会作为正式版本使用。",
            "删除版本会同步删除该版本下的接口和参数，操作前需要确认。",
        ],
    },
    {
        "level": 1,
        "title": "3. 导入原规格书",
        "paragraphs": [
            "导入原规格书用于上传已有 Word 通讯规格书。系统会保存原始模板，并尽量识别文档中的接口基础信息，后续导出 Word 时会基于该模板生成新版规格书。",
        ],
        "image": ("02-import-spec.png", "图 2 导入原规格书页面"),
        "steps": [
            "点击顶部或首页中的“导入原规格书”。",
            "选择 .docx 格式的规格书文件。",
            "点击上传，等待系统完成识别。",
            "导入完成后，返回版本管理页面确认版本是否正确。",
        ],
        "notes": [
            "当前仅支持 .docx 文件。",
            "如果原文档格式异常，接口识别可能不完整，需要在接口管理页面手动补充。",
        ],
    },
    {
        "level": 1,
        "title": "4. 接口管理工作台",
        "paragraphs": [
            "接口管理工作台用于查看当前版本下的所有接口。列表支持按方向和状态筛选，点击接口编号可以进入接口详情页。",
        ],
        "image": ("03-interface-list.png", "图 3 接口管理工作台"),
        "steps": [
            "从版本管理页点击“进入管理”。",
            "通过“全部 / EQP -> EAP / EAP -> EQP”筛选接口方向。",
            "通过状态筛选查看草稿或正式接口。",
            "点击接口编号进入接口详情，维护基础信息、参数和日志范例。",
        ],
        "notes": [
            "接口方向决定调用方和提供方的默认理解。",
            "导出 Word 或 PDF 成功后，对应版本接口会更新为正式状态。",
        ],
    },
    {
        "level": 1,
        "title": "5. 新增接口",
        "paragraphs": [
            "新增接口页面用于一次性录入接口基础信息、请求参数、响应参数和自定义子节点。页面支持按行维护参数，适合新增接口时批量录入字段。",
        ],
        "image": ("04-new-interface.png", "图 4 新增接口页面"),
        "steps": [
            "进入目标规格书版本，点击“新增接口”。",
            "填写接口编号、接口名称、接口方向、API 名称、调用方、提供方等基础信息。",
            "在请求参数区域按行新增字段。",
            "在响应参数区域按行新增字段。",
            "如字段类型为自定义对象或 List<自定义对象>，在下方维护对应自定义类型的子节点参数。",
            "确认无误后保存。",
        ],
        "notes": [
            "接口编号建议保持 EAP-EQP-xxx 或 EQP-EAP-xxx 格式。",
            "字段类型建议使用系统下拉选项，例如 string、int、float、double、bool、DateTime、object、array。",
            "字段说明较长时页面会自动换行显示，避免内容被截断。",
        ],
    },
    {
        "level": 1,
        "title": "6. 接口详情与参数维护",
        "paragraphs": [
            "接口详情页用于维护单个接口的完整内容。页面上方是接口基础信息，下方分别维护请求参数、响应参数和日志范例。",
        ],
        "image": ("05-interface-detail.png", "图 5 接口详情页面"),
        "steps": [
            "进入接口详情后，先确认接口基础信息是否正确。",
            "如需修改接口名称、API 名称、调用方、提供方或接口描述，编辑后保存。",
            "在请求参数或响应参数区域新增、修改、删除字段。",
            "字段调整后，重新生成或检查日志范例。",
        ],
        "notes": [
            "不要直接在已导出的旧版本上维护新需求，建议基于最新版本导出新版本后再维护。",
            "修改参数后要检查日志范例是否同步体现。",
        ],
    },
    {
        "level": 2,
        "title": "6.1 请求参数、响应参数与自定义子节点",
        "paragraphs": [
            "当某个字段是自定义类型，例如 StackUp 或 List<StackUp>，需要在对应参数区域下方维护自定义类型的子节点参数。导出 Word、PDF 以及生成日志范例时，系统会同步输出这些子节点。",
        ],
        "image": ("06-parameter-maintenance.png", "图 6 参数维护与自定义子节点区域"),
        "steps": [
            "在参数行中选择自定义类型，并填写自定义类型名称。",
            "在自定义类型模块中新增该类型的字段。",
            "确认子节点字段名称、类型、示例值和说明。",
            "保存后检查参数编号和日志范例。",
        ],
        "notes": [
            "同一个自定义类型可以包含多个子节点字段。",
            "如果请求参数和响应参数都有自定义类型，需要分别维护。",
        ],
    },
    {
        "level": 1,
        "title": "7. 日志范例生成",
        "paragraphs": [
            "日志范例用于给开发和测试人员快速确认接口请求、响应结构。系统会根据接口参数生成请求日志和响应 JSON 示例。",
        ],
        "image": ("07-log-examples.png", "图 7 日志范例区域"),
        "steps": [
            "完成接口基础信息和参数维护。",
            "进入接口详情页底部查看请求日志范例和响应日志范例。",
            "检查 RequestId、DateTime、字段类型和自定义子节点是否符合预期。",
            "如参数变更，保存后重新检查日志范例。",
        ],
        "notes": [
            "新增接口中的 RequestId 会按当前时间生成 17 位唯一标识格式。",
            "自定义子节点参数应同步体现在日志范例中。",
        ],
    },
    {
        "level": 1,
        "title": "8. 导出规格书",
        "paragraphs": [
            "导出中心支持导出 Markdown、Word、PDF、Word + PDF。导出 Word 或 PDF 成功后，系统会记录导出版本，并将对应版本接口状态更新为正式。",
        ],
        "image": ("08-export-center.png", "图 8 导出中心"),
        "steps": [
            "进入需要导出的规格书版本。",
            "点击“导出 Word/PDF”或版本卡片中的“导出”。",
            "选择导出格式。",
            "填写目标版本、修改人姓名和修改内容。",
            "选择或确认保存路径。",
            "点击导出，等待系统生成文件。",
        ],
        "notes": [
            "如果目标版本与当前版本一致，导出文件会归档到当前版本。",
            "如果目标版本是新版本，系统会复制当前版本接口到新版本。",
            "导出 Word 或 PDF 成功后，草稿接口会变为正式状态。",
        ],
    },
    {
        "level": 1,
        "title": "9. 服务器部署后的访问方式",
        "paragraphs": [
            "系统部署在服务器后，所有开发人员通过同一个服务器地址访问。数据库、导出文件和上传模板统一保存在服务器目录。",
        ],
        "bullets": [
            "服务器本机访问：http://127.0.0.1:8000。",
            "其他电脑访问：http://服务器IP:8000。",
            r"数据库位置：C:\EAPSystem\data\interface_manager.db。",
            r"导出文件位置：C:\EAPSystem\exports。",
            r"上传模板位置：C:\EAPSystem\uploads。",
        ],
        "notes": [
            "如果其他电脑无法访问，优先检查服务器防火墙是否放行 8000 端口。",
            "如果尚未安装为 Windows 服务，运行系统的 PowerShell 窗口不能关闭。",
        ],
    },
    {
        "level": 1,
        "title": "10. 常见问题",
        "qa": [
            ("本机能打开，其他电脑打不开怎么办？", "确认其他电脑使用的是服务器 IP，而不是 127.0.0.1；同时检查服务器防火墙是否放行 8000 端口。"),
            ("导入 Word 后接口识别不完整怎么办？", "先确认原 Word 格式是否规范；未识别的接口可以在对应版本中手动新增或补充。"),
            ("修改参数后日志范例没有变化怎么办？", "确认参数已保存，并刷新接口详情页；自定义类型需要同时维护子节点字段。"),
            ("导出后为什么接口状态变成正常？", "这是正式导出规则。Word 或 PDF 导出成功后，代表该版本已形成可发布文档，接口状态会更新为正式。"),
            ("数据应该如何备份？", r"重点备份 C:\EAPSystem\data\interface_manager.db，同时备份 uploads 和 exports 目录。"),
        ],
    },
]


def write_markdown() -> None:
    lines = [
        "# EAP接口规格书管理系统操作手册",
        "",
        f"更新日期：{TODAY}",
        "",
        "本文档用于指导开发人员和规格书维护人员使用 EAP接口规格书管理系统。后续系统功能更新时，请优先维护本 Markdown 文件，再同步生成 Word 版本。",
        "",
        "## 目录",
        "",
    ]
    for section in SECTIONS:
        indent = "  " if section["level"] == 2 else ""
        lines.append(f"{indent}- {section['title']}")
    lines.append("")

    for section in SECTIONS:
        heading = "#" * (section["level"] + 1)
        lines.append(f"{heading} {section['title']}")
        lines.append("")
        for text in section.get("paragraphs", []):
            lines.extend([text, ""])
        if "image" in section:
            image, caption = section["image"]
            lines.extend([f"![{caption}](images/manual/{image})", "", f"*{caption}*", ""])
        if "steps" in section:
            lines.extend(["**操作步骤**", ""])
            for index, text in enumerate(section["steps"], 1):
                lines.append(f"{index}. {text}")
            lines.append("")
        if "bullets" in section:
            lines.extend(["**要点说明**", ""])
            for text in section["bullets"]:
                lines.append(f"- {text}")
            lines.append("")
        if "notes" in section:
            lines.extend(["**注意事项**", ""])
            for text in section["notes"]:
                lines.append(f"- {text}")
            lines.append("")
        if "qa" in section:
            for question, answer in section["qa"]:
                lines.extend([f"**{question}**", "", answer, ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    accent = RGBColor(46, 116, 181)
    dark = RGBColor(31, 77, 120)
    muted = RGBColor(90, 103, 122)
    for style_name, size, color, before, after in [
        ("Heading 1", 16, accent, 14, 8),
        ("Heading 2", 13, accent, 10, 6),
        ("Heading 3", 12, dark, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EAP接口规格书管理系统操作手册")
    set_run_font(run, 22, RGBColor(11, 37, 69), True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"更新日期：{TODAY}    适用对象：开发人员 / 测试人员 / 规格书维护人员")
    set_run_font(run, 10, muted)

    intro = [
        ("系统用途", "集中维护 EAP-EQP Web API 通讯规格书，支持导入、维护、日志范例生成和导出。"),
        ("推荐访问", "http://服务器IP:8000"),
        ("维护方式", "以 Markdown 为源稿，Word 为发布版；后续功能更新时同步维护。"),
    ]
    table = doc.add_table(rows=len(intro), cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for row, (label, value) in zip(table.rows, intro):
        row.cells[0].width = Inches(1.25)
        row.cells[1].width = Inches(5.0)
        row.cells[0].text = label
        row.cells[1].text = value
        for cell_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                for cell_run in paragraph.runs:
                    set_run_font(cell_run, 9.5, dark if cell_index == 0 else None, cell_index == 0)

    doc.add_page_break()
    doc.add_heading("目录", level=1)
    for section_data in SECTIONS:
        paragraph = doc.add_paragraph()
        if section_data["level"] == 2:
            paragraph.paragraph_format.left_indent = Inches(0.25)
        run = paragraph.add_run(section_data["title"])
        set_run_font(run, 10.5)

    def add_label(text: str) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        set_run_font(run, 10.5, dark, True)

    def add_steps(items: list[str]) -> None:
        for item in items:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            run = paragraph.add_run(item)
            set_run_font(run, 10)

    def add_bullets(items: list[str]) -> None:
        for item in items:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.35)
            paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            run = paragraph.add_run(item)
            set_run_font(run, 10)

    def add_image(image_name: str, caption: str) -> None:
        image_path = IMAGE_DIR / image_name
        if not image_path.exists():
            return
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(6.25))
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, 9, muted)

    for section_data in SECTIONS:
        doc.add_heading(section_data["title"], level=section_data["level"])
        for text in section_data.get("paragraphs", []):
            paragraph = doc.add_paragraph(text)
            for run in paragraph.runs:
                set_run_font(run, 10.5)
        if "image" in section_data:
            add_image(*section_data["image"])
        if "steps" in section_data:
            add_label("操作步骤")
            add_steps(section_data["steps"])
        if "bullets" in section_data:
            add_label("要点说明")
            add_bullets(section_data["bullets"])
        if "notes" in section_data:
            add_label("注意事项")
            add_bullets(section_data["notes"])
        if "qa" in section_data:
            for question, answer in section_data["qa"]:
                add_label(question)
                paragraph = doc.add_paragraph(answer)
                for run in paragraph.runs:
                    set_run_font(run, 10.5)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("EAP接口规格书管理系统操作手册")
    set_run_font(run, 8, muted)

    doc.save(DOCX_PATH)


def main() -> None:
    DOCS_DIR.mkdir(exist_ok=True)
    write_markdown()
    build_docx()
    print(MD_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
