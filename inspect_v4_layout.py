from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from lxml import etree

DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")
doc = Document(str(DOCX))

print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt or i < 20:
        if "Version" in txt or "EAP-EQP" in txt or "PAGEBREAK" in txt:
            print("p", i, repr(txt), p.style.name)

with ZipFile(DOCX) as z:
    names = z.namelist()
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    doc_xml = etree.fromstring(z.read("word/document.xml"))
    brs = doc_xml.xpath(".//w:br[@w:type='page']", namespaces=ns)
    sects = doc_xml.xpath(".//w:sectPr", namespaces=ns)
    print("page_breaks", len(brs), "sectPr", len(sects))
    for name in names:
        if name.startswith("word/header") and name.endswith(".xml"):
            text = "".join(etree.fromstring(z.read(name)).xpath(".//w:t/text()", namespaces=ns))
            if text.strip():
                print(name, repr(text[:200]))
