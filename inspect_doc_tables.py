from docx import Document

doc = Document(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v3.8.docx")
print("tables", len(doc.tables))
for i, table in enumerate(doc.tables[:5]):
    print("table", i, "rows", len(table.rows), "cols", len(table.columns))
    for row in table.rows[:4]:
        print([cell.text for cell in row.cells])
