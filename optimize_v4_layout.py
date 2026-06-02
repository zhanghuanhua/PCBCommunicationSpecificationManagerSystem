from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOCX = Path(r"E:\MulTek\02设计文档\EAP通讯规格书\WebAPI\超毅项目Web API通讯规格书 v4.0.docx")


def set_run_font(run, font_name: str, size_pt: float, bold: bool | None = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = r_pr._add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{key}"), font_name)


def set_paragraph_compact(paragraph, font_name: str, size_pt: float, bold: bool | None = None) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    if not paragraph.runs and paragraph.text:
        paragraph.add_run("")
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt, bold)


def set_cell_margins(cell, top=40, start=60, bottom=40, end=60) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width_and_grid(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def is_example_table(table) -> bool:
    return "日志范例" in "\n".join(cell.text for row in table.rows for cell in row.cells)


def clean_label_text(text: str) -> str:
    text = re.sub(r"返回值列表表+", "返回值列表", text)
    text = re.sub(r"请求参数列表表+", "请求参数列表", text)
    text = re.sub(r"参数列表表+", "参数列表", text)
    text = re.sub(r"返回值列表\s*\|\s*返回值列表(?:\s*\|\s*返回值列表)+", "返回值列表", text)
    return text


def compact_json_text(text: str) -> str:
    if "{" not in text:
        return text
    normalized = text.replace("\u00a0", " ")
    normalized = normalized.replace("REST:POST ", "REST:POST ")
    lines = [line.rstrip() for line in normalized.splitlines() if line.strip()]
    prefix = []
    json_lines = []
    in_json = False
    for line in lines:
        if line.lstrip().startswith("{"):
            in_json = True
        if in_json:
            json_lines.append(line.strip())
        else:
            prefix.append(line.strip())
    if not json_lines:
        return "\n".join(lines)
    raw_json = "\n".join(json_lines)
    try:
        parsed = json.loads(raw_json)
        pretty = json.dumps(parsed, ensure_ascii=False, indent=2)
        return "\n".join(prefix + pretty.splitlines())
    except Exception:
        return "\n".join(lines)


def optimize_tables(doc: Document) -> None:
    for table in doc.tables:
        table.autofit = False
        example = is_example_table(table)
        if example:
            widths = [1250, 850, 7260] if len(table.columns) == 3 else [9360 // len(table.columns)] * len(table.columns)
            font, size = "Cambria", 7
        else:
            if len(table.columns) == 4:
                widths = [900, 1800, 1200, 5460]
            elif len(table.columns) == 5:
                widths = [700, 1250, 1250, 1250, 4910]
            elif len(table.columns) == 3:
                widths = [1100, 1800, 6460]
            elif len(table.columns) == 2:
                widths = [1500, 7860]
            else:
                widths = [max(800, 9360 // len(table.columns))] * len(table.columns)
            font, size = "Microsoft YaHei", 8
        set_table_width_and_grid(table, widths)
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            for h in tr_pr.findall(qn("w:trHeight")):
                tr_pr.remove(h)
            for cell in row.cells:
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER if not example else WD_CELL_VERTICAL_ALIGNMENT.TOP
                if example and ("{" in cell.text or "REST:POST" in cell.text):
                    cell.text = compact_json_text(cell.text)
                elif not example and "表表" in cell.text:
                    cell.text = clean_label_text(cell.text)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_paragraph_compact(paragraph, font, size)
        if example:
            for row in table.rows:
                if row.cells:
                    for p in row.cells[0].paragraphs:
                        set_paragraph_compact(p, "Cambria", 7, bold=True)


def update_headers(doc: Document) -> None:
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                if "Version:" in paragraph.text:
                    paragraph.text = paragraph.text.replace("Version: 3.8", "Version: 4.0")
                for run in paragraph.runs:
                    if "Version: 3.8" in run.text:
                        run.text = run.text.replace("Version: 3.8", "Version: 4.0")


def remove_empty_page_break_paragraphs(docx_path: Path) -> None:
    tmp = docx_path.with_suffix(".tmp.docx")
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                body = root.find("w:body", namespaces=ns)
                blocks = list(body)
                for p in blocks:
                    text = "".join(p.xpath(".//w:t/text()", namespaces=ns)).strip()
                    has_page_break = bool(p.xpath(".//w:br[@w:type='page']", namespaces=ns))
                    if has_page_break and not text:
                        # Keep intentional breaks before major direction sections, remove repeated empty breaks.
                        prev_text = "".join(blocks[blocks.index(p) - 1].xpath(".//w:t/text()", namespaces=ns)).strip() if blocks.index(p) > 0 else ""
                        next_text = "".join(blocks[blocks.index(p) + 1].xpath(".//w:t/text()", namespaces=ns)).strip() if blocks.index(p) + 1 < len(blocks) else ""
                        if not next_text or prev_text.startswith("日志范例"):
                            body.remove(p)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
            elif item.filename.startswith("word/header") and item.filename.endswith(".xml"):
                text = data.decode("utf-8")
                text = text.replace("Version: 3.8", "Version: 4.0")
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(docx_path)


def main() -> None:
    doc = Document(str(DOCX))
    update_headers(doc)
    optimize_tables(doc)
    doc.save(str(DOCX))
    remove_empty_page_break_paragraphs(DOCX)
    print(f"optimized={DOCX}")


if __name__ == "__main__":
    from lxml import etree

    main()
